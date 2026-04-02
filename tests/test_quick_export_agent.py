# -*- coding: utf-8 -*-
from datetime import datetime, timezone
from importlib import import_module
from pathlib import Path
from types import SimpleNamespace
import sys
import types
import pytest

from docx import Document

sys.modules.setdefault('pypdf', types.SimpleNamespace(PdfReader=None))
sys.modules.setdefault('pdfplumber', types.SimpleNamespace(open=None))
agents_pkg = types.ModuleType('src.agents')
agents_pkg.__path__ = [str(Path(__file__).resolve().parents[1] / 'src' / 'agents')]
sys.modules.setdefault('src.agents', agents_pkg)

from src.utils.xml_security import parse_xml_safely

QuickExportAgent = import_module('src.agents.quick_export_agent').QuickExportAgent


class _DummyLLM:
    def call(self, *args, **kwargs):
        return {}


class _DummyDB:
    def add(self, *_args, **_kwargs):
        return None

    def commit(self):
        return None

    def refresh(self, _obj):
        return None


def _make_agent(tmp_path: Path) -> QuickExportAgent:
    return QuickExportAgent(
        llm_gateway=_DummyLLM(),
        db_session=_DummyDB(),
        config={'export_dir': str(tmp_path)},
    )


def test_export_docx_preserves_original_bytes(tmp_path: Path):
    original_path = tmp_path / 'source.docx'
    doc = Document()
    doc.add_heading('Договор поставки', 0)
    doc.add_paragraph('Поставщик обязуется поставить товар.')
    doc.save(original_path)

    contract = SimpleNamespace(
        id='contract-1',
        file_name='source.docx',
        file_path=str(original_path),
        analysis_results=[],
        meta_info={},
        upload_date=datetime.now(timezone.utc),
        document_type='contract',
        contract_type='supply',
        status='completed',
        risk_level=None,
    )

    agent = _make_agent(tmp_path)
    export_path = agent._export_docx(contract, 'exported', include_analysis=False)

    assert Path(export_path).exists()
    assert Path(export_path).read_bytes() == original_path.read_bytes()


def test_export_xml_returns_canonical_valid_xml(tmp_path: Path):
    original_path = tmp_path / 'source.docx'
    doc = Document()
    doc.add_heading('Договор поставки', 0)
    doc.add_paragraph('1. Предмет договора')
    doc.add_paragraph('Поставщик обязуется поставить зерно.')
    doc.save(original_path)

    contract = SimpleNamespace(
        id='contract-2',
        file_name='source.docx',
        file_path=str(original_path),
        analysis_results=[],
        meta_info={},
        upload_date=datetime.now(timezone.utc),
        document_type='contract',
        contract_type='supply',
        status='completed',
        risk_level=None,
    )

    agent = _make_agent(tmp_path)
    export_path = agent._export_xml(contract, 'canonical', include_analysis=False)

    xml_text = Path(export_path).read_text(encoding='utf-8')
    root = parse_xml_safely(xml_text)

    assert root.tag == 'contract'
    assert root.find('.//clauses') is not None


def test_cross_format_export_requires_explicit_lossy_confirmation(tmp_path: Path):
    source_path = tmp_path / 'source.pdf'
    source_path.write_bytes(b'%PDF-1.4 placeholder')

    contract = SimpleNamespace(
        id='contract-3',
        file_name='source.pdf',
        file_path=str(source_path),
        analysis_results=[],
        meta_info={
            'xml': '<contract><title>Договор</title><clauses><clause><title>1. Предмет</title><paragraph>Текст договора.</paragraph></clause></clauses></contract>'
        },
        upload_date=datetime.now(timezone.utc),
        document_type='contract',
        contract_type='supply',
        status='completed',
        risk_level=None,
    )

    agent = _make_agent(tmp_path)

    with pytest.raises(ValueError, match='подтвердите предупреждение'):
        agent._export_docx(contract, 'converted', include_analysis=False, allow_lossy_conversion=False)


def test_cross_format_export_generates_best_effort_docx_after_confirmation(tmp_path: Path):
    source_path = tmp_path / 'source.pdf'
    source_path.write_bytes(b'%PDF-1.4 placeholder')

    contract = SimpleNamespace(
        id='contract-4',
        file_name='source.pdf',
        file_path=str(source_path),
        analysis_results=[],
        meta_info={
            'xml': '<contract><metadata><title>Без названия</title></metadata><clauses><clause type="preamble"><title>Преамбула</title><content><paragraph>ДОГОВОР ПОСТАВКИ</paragraph></content></clause><clause><title>1. Предмет</title><content><paragraph>Поставщик обязуется поставить зерно.</paragraph></content></clause></clauses><tables><table><row><cell>Продавец: ООО Ромашка</cell><cell>Покупатель: ООО Василек</cell></row></table></tables></contract>'
        },
        upload_date=datetime.now(timezone.utc),
        document_type='contract',
        contract_type='supply',
        status='completed',
        risk_level=None,
    )

    agent = _make_agent(tmp_path)
    export_path = agent._export_docx(contract, 'converted', include_analysis=False, allow_lossy_conversion=True)

    exported_doc = Document(export_path)
    exported_text = '\n'.join(p.text for p in exported_doc.paragraphs if p.text.strip())

    assert Path(export_path).exists()
    assert 'ДОГОВОР ПОСТАВКИ' in exported_text
    assert 'Поставщик обязуется поставить зерно.' in exported_text
    assert 'converted' not in exported_text
    assert len(exported_doc.tables) == 1
    assert exported_doc.tables[0].cell(0, 0).text == 'Продавец: ООО Ромашка'


def test_cross_format_export_generates_best_effort_pdf_after_confirmation(tmp_path: Path):
    pytest.importorskip('reportlab')

    source_path = tmp_path / 'source.docx'
    doc = Document()
    doc.add_paragraph('Исходный DOCX')
    doc.save(source_path)

    contract = SimpleNamespace(
        id='contract-5',
        file_name='source.docx',
        file_path=str(source_path),
        analysis_results=[],
        meta_info={
            'xml': '<contract><metadata><title>Без названия</title></metadata><clauses><clause type="preamble"><title>Преамбула</title><content><paragraph>ДОГОВОР УСЛУГ</paragraph></content></clause><clause><title>1. Предмет</title><content><paragraph>Исполнитель оказывает услуги.</paragraph></content></clause></clauses><tables><table><row><cell>Исполнитель: ООО Альфа</cell><cell>Заказчик: ООО Бета</cell></row></table></tables></contract>'
        },
        upload_date=datetime.now(timezone.utc),
        document_type='contract',
        contract_type='service',
        status='completed',
        risk_level=None,
    )

    agent = _make_agent(tmp_path)
    export_path = agent._export_pdf(contract, 'converted_pdf', include_analysis=False, allow_lossy_conversion=True)

    from pypdf import PdfReader
    pdf_text = '\n'.join((page.extract_text() or '') for page in PdfReader(export_path).pages)

    assert Path(export_path).exists()
    assert Path(export_path).stat().st_size > 0
    assert 'ДОГОВОР УСЛУГ' in pdf_text
    assert 'Исполнитель оказывает услуги.' in pdf_text
    assert 'Исполнитель: ООО Альфа' in pdf_text
