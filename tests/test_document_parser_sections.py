# -*- coding: utf-8 -*-
from pathlib import Path
import xml.etree.ElementTree as ET

from docx import Document

from src.services.document_parser import DocumentParser


def test_docx_parser_splits_numbered_sections_and_preserves_tables(tmp_path: Path):
    docx_path = tmp_path / "supply_contract.docx"
    doc = Document()
    doc.add_paragraph("ДОГОВОР ПОСТАВКИ № 1")
    doc.add_paragraph("ООО «Поставщик» и ООО «Покупатель» заключили договор.")
    doc.add_paragraph("1. Предмет договора")
    doc.add_paragraph("1.1. Поставщик поставляет зерно.")
    doc.add_paragraph("2. Цена договора")
    doc.add_paragraph("2.1. Цена определяется в счете.")
    doc.add_paragraph("8. Реквизиты и подписи сторон")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Продавец: ООО «Поставщик»"
    table.rows[0].cells[1].text = "Покупатель: ООО «Покупатель»"
    table.rows[1].cells[0].text = "ИНН 1234567890"
    table.rows[1].cells[1].text = "ИНН 0987654321"
    doc.save(docx_path)

    parser = DocumentParser()
    xml_content = parser.parse(str(docx_path))
    root = ET.fromstring(xml_content)

    clause_titles = [(clause.findtext("title") or "").strip() for clause in root.findall(".//clauses/clause")]
    assert clause_titles[:4] == [
        "Преамбула",
        "1. Предмет договора",
        "2. Цена договора",
        "8. Реквизиты и подписи сторон",
    ]

    cells = ["".join(cell.itertext()).strip() for cell in root.findall(".//tables/table/row/cell")]
    assert "Продавец: ООО «Поставщик»" in cells
    assert "ИНН 0987654321" in cells
