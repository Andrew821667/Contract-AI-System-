# -*- coding: utf-8 -*-
"""
Contract Graph Parser

Парсинг договоров (DOCX, PDF, TXT, XML) в нормализованное дерево GraphNode.
Приоритет определения структуры:
  1. Нумерация (1. → 1.1 → 1.1.1) — основной сигнал
  2. DOCX styles (Heading1/2/3)
  3. Паттерны текста (регулярки для преамбулы, приложений, таблиц)
"""
from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
from lxml import etree

from .base_parser import BaseDocumentGraphParser, ParsedNode, ParseResult
from ..enums import LayerType, NodeType, ParseStatus


# ──────────────────────────────────────────────
# Regex patterns для определения структуры
# ──────────────────────────────────────────────

# Нумерованный пункт: "1.", "1.1", "1.1.", "1.1.1", "1.1.1."
# Для однокомпонентных номеров точка обязательна: "1. Текст"
# Для многокомпонентных (1.1, 2.3.1) точка опциональна: "1.1 Текст" или "1.1. Текст"
RE_NUMBERED = re.compile(
    r'^(\d+\.\d+(?:\.\d+)*)\.?\s+(.*)'  # 1.1, 1.1.1, ... — точка опциональна
    r'|^(\d+)\.\s+(.*)',                  # 1. 2. 3. — точка обязательна
    re.MULTILINE
)

# Раздел: "Раздел 1", "РАЗДЕЛ I", "Глава 2"
RE_SECTION_HEADER = re.compile(
    r'^(?:Раздел|РАЗДЕЛ|Глава|ГЛАВА)\s+(\d+|[IVXLC]+)[\.:]?\s*(.*)',
    re.IGNORECASE
)

# Приложение: "Приложение №1", "Приложение 1", "ПРИЛОЖЕНИЕ N 1"
RE_APPENDIX = re.compile(
    r'^(?:Приложение|ПРИЛОЖЕНИЕ)\s*[№NnО]?\s*(\d+)',
    re.IGNORECASE
)

# Таблица или список
RE_TABLE_HEADER = re.compile(r'^(?:Таблица|ТАБЛИЦА)\s*[№N]?\s*(\d+)', re.IGNORECASE)

# Блок подписей (эвристика: строки с "подпись", "М.П.", "___")
# НЕ включает "Поставщик...Покупатель" без ____ — это может быть преамбула
RE_SIGNATURE = re.compile(
    r'(?:подпис|М\.?\s*П\.|_{5,})',
    re.IGNORECASE
)

# Преамбула: начинается с организации или даты
RE_PREAMBLE = re.compile(
    r'^(?:.*(?:ООО|ОАО|ПАО|АО|ЗАО|ИП|ФГУП)\s+"[^"]+"|.*именуем)',
    re.IGNORECASE
)

# Классификация раздела по заголовку
SECTION_TYPE_KEYWORDS: Dict[str, List[str]] = {
    'subject': ['предмет', 'subject'],
    'financial': ['цена', 'расчет', 'оплата', 'стоимость', 'вознаграждение', 'payment', 'price'],
    'terms': ['срок', 'период', 'действи', 'term', 'deadline'],
    'liability': ['ответственность', 'штраф', 'неустойк', 'пени', 'liability'],
    'dispute': ['спор', 'арбитраж', 'претенз', 'dispute'],
    'confidentiality': ['конфиденциальн', 'тайн', 'confidential'],
    'force_majeure': ['форс-мажор', 'непреодол', 'force majeure'],
    'termination': ['расторж', 'прекращ', 'termination'],
    'warranty': ['гарант', 'качеств', 'warranty'],
    'definitions': ['определен', 'термин', 'понятия', 'definition'],
    'delivery': ['поставк', 'доставк', 'отгрузк', 'delivery'],
    'acceptance': ['приёмк', 'приемк', 'acceptance'],
}


def _classify_section(title: str) -> str:
    """Классификация раздела по заголовку."""
    title_lower = title.lower()
    for section_type, keywords in SECTION_TYPE_KEYWORDS.items():
        if any(kw in title_lower for kw in keywords):
            return section_type
    return "general"


def _numbering_level(number: str) -> int:
    """Уровень вложенности по нумерации: '1' → 0, '1.1' → 1, '1.1.1' → 2."""
    return number.count('.')


def _numbering_parent(number: str) -> Optional[str]:
    """Родительский номер: '1.1.1' → '1.1', '1.1' → '1', '1' → None."""
    parts = number.split('.')
    if len(parts) <= 1:
        return None
    return '.'.join(parts[:-1])


# ──────────────────────────────────────────────
# ContractGraphParser
# ──────────────────────────────────────────────

class ContractGraphParser(BaseDocumentGraphParser):
    """
    Парсер договоров в нормализованное дерево.

    Три режима работы:
    1. parse_file(path) — полный pipeline (DOCX/PDF/TXT)
    2. parse_text(text) — парсинг plain text
    3. parse_xml(xml) — парсинг XML от существующего DocumentParser
    """

    def parse_file(self, file_path: str) -> ParseResult:
        """Распарсить файл договора."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = Path(file_path).suffix.lower()
        file_name = Path(file_path).name

        if ext == '.docx':
            return self._parse_docx(file_path)
        elif ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
            result = self.parse_text(text, title=file_name)
            result.source_format = 'txt'
            return result
        else:
            raise ValueError(f"Unsupported format: {ext}")

    def parse_text(self, text: str, title: str = "Без названия") -> ParseResult:
        """Распарсить plain text договора в дерево."""
        root = ParsedNode(
            node_type=NodeType.DOCUMENT,
            text=title,
            title=title,
            level=0,
            position=0,
        )

        lines = text.split('\n')
        lines = [ln.strip() for ln in lines if ln.strip()]

        if not lines:
            return ParseResult(
                root=root, layer=LayerType.CONTRACT, title=title,
                parse_status=ParseStatus.FAILED,
                parse_errors=["Empty document"],
            )

        # Pass 1: классифицируем каждую строку
        classified = self._classify_lines(lines)

        # Pass 2: строим дерево по нумерации и классификации
        self._build_tree(root, classified)

        parse_status = ParseStatus.FULLY_PARSED if root.children else ParseStatus.PARTIAL_PARSE
        errors = [] if root.children else ["No structure detected"]

        return ParseResult(
            root=root,
            layer=LayerType.CONTRACT,
            title=title,
            source_format='txt',
            parse_status=parse_status,
            parse_errors=errors,
        )

    def parse_xml(self, xml_content: str, title: str = "Без названия") -> ParseResult:
        """
        Парсинг XML от существующего DocumentParser.

        Конвертирует плоскую структуру <clauses><clause> в настоящее дерево
        с вложенностью по нумерации.
        """
        try:
            # Безопасный парсинг XML
            parser = etree.XMLParser(resolve_entities=False, no_network=True)
            tree = etree.fromstring(xml_content.encode('utf-8') if isinstance(xml_content, str) else xml_content,
                                    parser=parser)
        except etree.XMLSyntaxError as e:
            return ParseResult(
                root=ParsedNode(node_type=NodeType.DOCUMENT, text=title, title=title),
                layer=LayerType.CONTRACT, title=title,
                parse_status=ParseStatus.FAILED,
                parse_errors=[f"XML parse error: {e}"],
            )

        root = ParsedNode(
            node_type=NodeType.DOCUMENT,
            text=title,
            title=title,
            level=0,
            position=0,
        )

        # Извлекаем метаданные
        metadata = {}
        meta_elem = tree.find('.//metadata')
        if meta_elem is not None:
            for child in meta_elem:
                if child.text:
                    metadata[child.tag] = child.text
        root.metadata = metadata

        # Преамбула: parties + terms
        preamble_parts = []
        parties_elem = tree.find('.//parties')
        if parties_elem is not None:
            for party in parties_elem.findall('party'):
                name = party.findtext('name', '')
                role = party.get('role', '')
                if name:
                    preamble_parts.append(f"{role}: {name}")

        if preamble_parts:
            preamble = root.add_child(ParsedNode(
                node_type=NodeType.PREAMBLE,
                text='; '.join(preamble_parts),
                title="Преамбула",
            ))

        # Clauses → дерево по нумерации
        clauses_elem = tree.find('.//clauses')
        if clauses_elem is not None:
            clause_elements = clauses_elem.findall('clause')
            self._xml_clauses_to_tree(root, clause_elements)

        # Таблицы
        tables_elem = tree.find('.//tables')
        if tables_elem is not None:
            for idx, table_elem in enumerate(tables_elem.findall('table')):
                rows_text = []
                for row in table_elem.findall('row'):
                    cells = [c.text or '' for c in row.findall('cell')]
                    rows_text.append(' | '.join(cells))
                if rows_text:
                    root.add_child(ParsedNode(
                        node_type=NodeType.TABLE,
                        text='\n'.join(rows_text),
                        title=f"Таблица {idx + 1}",
                        number=str(idx + 1),
                        metadata={"table_id": table_elem.get('id', str(idx + 1))},
                    ))

        parse_status = ParseStatus.FULLY_PARSED if root.children else ParseStatus.PARTIAL_PARSE

        return ParseResult(
            root=root,
            layer=LayerType.CONTRACT,
            title=metadata.get('title', title),
            source_format='xml',
            parse_status=parse_status,
            metadata=metadata,
        )

    # ──────────────────────────────────────────
    # DOCX parsing
    # ──────────────────────────────────────────

    def _parse_docx(self, file_path: str) -> ParseResult:
        """Парсинг DOCX с учётом стилей и нумерации."""
        try:
            from docx import Document
        except ImportError:
            return ParseResult(
                root=ParsedNode(node_type=NodeType.DOCUMENT, text=Path(file_path).name),
                layer=LayerType.CONTRACT, title=Path(file_path).name,
                parse_status=ParseStatus.FAILED,
                parse_errors=["python-docx not installed"],
            )

        doc = Document(file_path)
        file_name = Path(file_path).name

        # Собираем параграфы с метаданными стилей
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ''
            is_heading = style_name.startswith('Heading')
            heading_level = 0
            if is_heading:
                try:
                    heading_level = int(style_name.replace('Heading', '').replace(' ', ''))
                except ValueError:
                    heading_level = 1
            paragraphs.append({
                'text': text,
                'style': style_name,
                'is_heading': is_heading,
                'heading_level': heading_level,
            })

        # Таблицы
        tables = []
        for idx, table in enumerate(doc.tables):
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(' | '.join(cells))
            if rows_text:
                tables.append({'text': '\n'.join(rows_text), 'index': idx + 1})

        # Строим дерево
        root = ParsedNode(
            node_type=NodeType.DOCUMENT,
            text=file_name,
            title=doc.core_properties.title or file_name,
            metadata={
                'author': doc.core_properties.author or '',
                'file_name': file_name,
            },
        )

        # Конвертируем в классифицированные строки и строим дерево
        classified = self._classify_docx_paragraphs(paragraphs)
        self._build_tree(root, classified)

        # Добавляем таблицы
        for tbl in tables:
            root.add_child(ParsedNode(
                node_type=NodeType.TABLE,
                text=tbl['text'],
                title=f"Таблица {tbl['index']}",
                number=str(tbl['index']),
            ))

        parse_status = ParseStatus.FULLY_PARSED if root.children else ParseStatus.PARTIAL_PARSE

        return ParseResult(
            root=root,
            layer=LayerType.CONTRACT,
            title=doc.core_properties.title or file_name,
            source_format='docx',
            parse_status=parse_status,
            metadata={'author': doc.core_properties.author or ''},
        )

    # ──────────────────────────────────────────
    # PDF parsing
    # ──────────────────────────────────────────

    def _parse_pdf(self, file_path: str) -> ParseResult:
        """Парсинг PDF: извлекаем текст и парсим как plain text."""
        text_content = []
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
        except Exception:
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            except Exception as e:
                return ParseResult(
                    root=ParsedNode(node_type=NodeType.DOCUMENT, text=Path(file_path).name),
                    layer=LayerType.CONTRACT, title=Path(file_path).name,
                    parse_status=ParseStatus.FAILED,
                    parse_errors=[f"PDF extraction failed: {e}"],
                )

        full_text = '\n\n'.join(text_content)
        file_name = Path(file_path).name

        result = self.parse_text(full_text, title=file_name)
        result.source_format = 'pdf'
        return result

    # ──────────────────────────────────────────
    # Internal: классификация строк
    # ──────────────────────────────────────────

    def _classify_lines(self, lines: List[str]) -> List[Dict[str, Any]]:
        """
        Классификация каждой строки: определение типа, номера, уровня.
        """
        classified = []
        preamble_ended = False

        for line in lines:
            entry: Dict[str, Any] = {'text': line, 'node_type': NodeType.PARAGRAPH,
                                      'number': None, 'title': None, 'level': None,
                                      'heading_level': 0}

            # 1. Приложение
            m = RE_APPENDIX.match(line)
            if m:
                entry['node_type'] = NodeType.APPENDIX
                entry['number'] = m.group(1)
                entry['title'] = line
                classified.append(entry)
                continue

            # 2. Нумерованный пункт (основной сигнал!)
            m = RE_NUMBERED.match(line)
            if m:
                preamble_ended = True
                # Две альтернативы: groups (1,2) для multi-component, (3,4) для single
                number = m.group(1) or m.group(3)
                text_after = m.group(2) or m.group(4) or ''
                lvl = _numbering_level(number)

                if lvl == 0:
                    entry['node_type'] = NodeType.SECTION
                    entry['title'] = text_after
                elif lvl == 1:
                    entry['node_type'] = NodeType.CLAUSE
                else:
                    entry['node_type'] = NodeType.SUBCLAUSE

                entry['number'] = number
                entry['text'] = text_after
                entry['level'] = lvl
                classified.append(entry)
                continue

            # 3. Заголовок раздела без нумерации
            m = RE_SECTION_HEADER.match(line)
            if m:
                preamble_ended = True
                entry['node_type'] = NodeType.SECTION
                entry['number'] = m.group(1)
                entry['title'] = m.group(2) or line
                entry['text'] = line
                classified.append(entry)
                continue

            # 4. Блок подписей
            if RE_SIGNATURE.search(line):
                entry['node_type'] = NodeType.SIGNATURE_BLOCK
                classified.append(entry)
                continue

            # 5. Преамбула (до первого нумерованного пункта)
            if not preamble_ended:
                entry['node_type'] = NodeType.PREAMBLE
                classified.append(entry)
                continue

            # 6. Обычный параграф
            classified.append(entry)

        return classified

    def _classify_docx_paragraphs(self, paragraphs: List[Dict]) -> List[Dict[str, Any]]:
        """Классификация DOCX-параграфов с учётом стилей."""
        classified = []
        preamble_ended = False

        for para in paragraphs:
            text = para['text']
            entry: Dict[str, Any] = {'text': text, 'node_type': NodeType.PARAGRAPH,
                                      'number': None, 'title': None, 'level': None,
                                      'heading_level': para.get('heading_level', 0)}

            # Heading стиль → section
            if para.get('is_heading'):
                preamble_ended = True
                entry['node_type'] = NodeType.SECTION
                entry['title'] = text
                # Попробуем извлечь номер из текста
                m = RE_NUMBERED.match(text)
                if m:
                    number = m.group(1) or m.group(3)
                    text_after = m.group(2) or m.group(4) or ''
                    entry['number'] = number
                    entry['text'] = text_after
                    entry['level'] = _numbering_level(number)
                classified.append(entry)
                continue

            # Нумерованный пункт
            m = RE_NUMBERED.match(text)
            if m:
                preamble_ended = True
                number = m.group(1) or m.group(3)
                text_after = m.group(2) or m.group(4) or ''
                lvl = _numbering_level(number)
                if lvl == 0:
                    entry['node_type'] = NodeType.SECTION
                    entry['title'] = text_after
                elif lvl == 1:
                    entry['node_type'] = NodeType.CLAUSE
                else:
                    entry['node_type'] = NodeType.SUBCLAUSE
                entry['number'] = number
                entry['text'] = text_after
                entry['level'] = lvl
                classified.append(entry)
                continue

            # Приложение
            m_app = RE_APPENDIX.match(text)
            if m_app:
                entry['node_type'] = NodeType.APPENDIX
                entry['number'] = m_app.group(1)
                entry['title'] = text
                classified.append(entry)
                continue

            # Подпись
            if RE_SIGNATURE.search(text):
                entry['node_type'] = NodeType.SIGNATURE_BLOCK
                classified.append(entry)
                continue

            # Преамбула
            if not preamble_ended:
                entry['node_type'] = NodeType.PREAMBLE
                classified.append(entry)
                continue

            classified.append(entry)

        return classified

    # ──────────────────────────────────────────
    # Internal: построение дерева
    # ──────────────────────────────────────────

    def _build_tree(self, root: ParsedNode, classified: List[Dict[str, Any]]):
        """
        Построение дерева из классифицированных строк.

        Логика:
        - Преамбула → дочерний узел root
        - Нумерованные пункты → дерево по нумерации (1 → 1.1 → 1.1.1)
        - Ненумерованные параграфы → присоединяются к последнему нумерованному
        - Приложения, таблицы, подписи → дочерние root
        """
        # Собираем преамбулу
        preamble_lines = []
        numbered_and_rest = []
        for entry in classified:
            if entry['node_type'] == NodeType.PREAMBLE:
                preamble_lines.append(entry['text'])
            else:
                numbered_and_rest.append(entry)

        if preamble_lines:
            root.add_child(ParsedNode(
                node_type=NodeType.PREAMBLE,
                text='\n'.join(preamble_lines),
                title="Преамбула",
            ))

        # Индекс нумерованных узлов: "1" → node, "1.1" → node
        number_index: Dict[str, ParsedNode] = {}
        last_numbered: Optional[ParsedNode] = None
        signature_lines: List[str] = []

        for entry in numbered_and_rest:
            ntype = entry['node_type']
            number = entry.get('number')

            if ntype == NodeType.SIGNATURE_BLOCK:
                signature_lines.append(entry['text'])
                continue

            if ntype == NodeType.APPENDIX:
                root.add_child(ParsedNode(
                    node_type=NodeType.APPENDIX,
                    text=entry['text'],
                    title=entry.get('title', entry['text']),
                    number=number,
                ))
                continue

            if number and ntype in (NodeType.SECTION, NodeType.CLAUSE, NodeType.SUBCLAUSE):
                node = ParsedNode(
                    node_type=ntype,
                    text=entry['text'],
                    title=entry.get('title'),
                    number=number,
                    metadata={'section_type': _classify_section(entry.get('title') or entry['text'])},
                )

                # Определяем родителя по нумерации
                parent_number = _numbering_parent(number)
                if parent_number and parent_number in number_index:
                    number_index[parent_number].add_child(node)
                else:
                    root.add_child(node)

                number_index[number] = node
                last_numbered = node
                continue

            # Heading без нумерации
            if ntype == NodeType.SECTION and not number:
                node = ParsedNode(
                    node_type=NodeType.SECTION,
                    text=entry['text'],
                    title=entry.get('title', entry['text']),
                )
                root.add_child(node)
                last_numbered = node
                continue

            # Обычный параграф → присоединяем к последнему нумерованному
            if ntype == NodeType.PARAGRAPH:
                if last_numbered:
                    last_numbered.add_child(ParsedNode(
                        node_type=NodeType.PARAGRAPH,
                        text=entry['text'],
                    ))
                else:
                    root.add_child(ParsedNode(
                        node_type=NodeType.PARAGRAPH,
                        text=entry['text'],
                    ))

        # Блок подписей
        if signature_lines:
            root.add_child(ParsedNode(
                node_type=NodeType.SIGNATURE_BLOCK,
                text='\n'.join(signature_lines),
                title="Подписи сторон",
            ))

    # ──────────────────────────────────────────
    # Internal: XML clauses → дерево
    # ──────────────────────────────────────────

    def _xml_clauses_to_tree(self, root: ParsedNode, clause_elements: list):
        """Конвертация плоских <clause> из DocumentParser XML в дерево по нумерации."""
        number_index: Dict[str, ParsedNode] = {}

        for clause_elem in clause_elements:
            title_elem = clause_elem.find('title')
            content_elem = clause_elem.find('content')

            title_text = title_elem.text.strip() if title_elem is not None and title_elem.text else ''
            clause_type = clause_elem.get('type', 'general')

            # Собираем текст
            if content_elem is not None:
                paragraphs = content_elem.findall('paragraph')
                text = '\n'.join([p.text for p in paragraphs if p.text])
            else:
                text = ''.join(clause_elem.itertext()).strip()

            if not text and not title_text:
                continue

            # Пытаемся извлечь номер из заголовка
            number = None
            display_title = title_text
            m = RE_NUMBERED.match(title_text)
            if m:
                number = m.group(1) or m.group(3)
                display_title = m.group(2) or m.group(4) or ''

            # Определяем тип узла
            if number:
                lvl = _numbering_level(number)
                if lvl == 0:
                    node_type = NodeType.SECTION
                elif lvl == 1:
                    node_type = NodeType.CLAUSE
                else:
                    node_type = NodeType.SUBCLAUSE
            else:
                node_type = NodeType.CLAUSE

            node = ParsedNode(
                node_type=node_type,
                text=text or display_title,
                title=display_title or None,
                number=number,
                metadata={
                    'section_type': clause_type,
                    'xml_id': clause_elem.get('id'),
                },
            )

            # Определяем родителя по нумерации
            if number:
                parent_number = _numbering_parent(number)
                if parent_number and parent_number in number_index:
                    number_index[parent_number].add_child(node)
                else:
                    root.add_child(node)
                number_index[number] = node
            else:
                root.add_child(node)
