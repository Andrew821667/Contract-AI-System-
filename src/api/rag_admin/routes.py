# -*- coding: utf-8 -*-
"""
RAG Admin API — управление базой знаний ChromaDB.

Endpoints:
  GET    /api/v1/rag/stats                     — статистика по коллекциям
  GET    /api/v1/rag/documents?collection=...  — список документов
  POST   /api/v1/rag/documents                 — загрузить файл
  DELETE /api/v1/rag/documents/{doc_id}        — удалить документ

Синглтоны ChromaDB/эмбеддинги — из src.services.admin_rag_retriever (общие с агентом).
"""
import collections
import hashlib
import io
import json
import re
import threading
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile, status
from loguru import logger
from pydantic import BaseModel

from src.api.dependencies import get_current_user
from src.models.auth_models import User
from src.services.admin_rag_retriever import (
    COLLECTION_LABELS,
    COLLECTIONS,
    get_collection as _get_collection_shared,
)

# БЕЗ prefix="/rag": роутер уже монтируется в main.py как prefix="/api/v1/rag".
# Свой префикс задваивал путь → реальные эндпоинты жили на /api/v1/rag/rag/*, а
# фронтенд (api.ts: /api/v1/rag/stats, /api/v1/rag/documents) бил в /api/v1/rag/*
# и получал 404 — RAG-админка в UI была нерабочей (L11). Двойной путь не ждёт
# никто, поэтому убираем префикс: серверные пути сходятся с клиентскими.
router = APIRouter(tags=["RAG Admin"])

RAG_MAX_FILE_SIZE = 25 * 1024 * 1024  # 25 МБ — лимит на загрузку документа в БЗ

# Stats cache (60s TTL) — stats query scans all collections, cache per request
_stats_cache: Optional[tuple] = None  # (StatsResponse, expires_at: float)
_STATS_TTL = 60.0

# ── Кеш числа документов ────────────────────────────────────────────────────
# doc_count = число уникальных doc_id среди метаданных ВСЕХ чанков коллекции.
# На проде это 1.3 млн чанков в 20-ГБ сторе: полный обход через chroma занимает
# ~60с (laws ~23с + case_law ~39с), прямой SQL по chroma.sqlite3 — ещё хуже
# (2м13с на холодном кеше). Таймаут axios во фронте — 30с, поэтому синхронный
# расчёт ГАРАНТИРОВАННО обрывался и админ-панель показывала пустоту.
# Решение: ответ отдаём мгновенно (chunk_count дёшев через count()), а doc_count
# берём из кеша, который пересчитывает фоновый поток. Кеш переживает рестарт
# (пишется на диск), иначе после каждого рестарта панель снова показывала бы нули.
_DOC_COUNT_TTL = 6 * 3600.0     # пересчёт не чаще раза в 6 часов
_DOC_COUNT_PAGE = 10000         # размер страницы при обходе метаданных
_DOC_SCAN_MAX_CHUNKS = 50000    # потолок обхода в /documents, чтобы не подвесить панель
_DOC_COUNT_FILE = Path("data/rag_doc_counts.json")
_doc_counts: Dict[str, int] = {}
_doc_counts_at: float = 0.0
_doc_count_lock = threading.Lock()
_doc_count_running = False

# Upload rate limiter: 10 uploads per user per minute (token bucket per user_id)
_upload_rl_lock = threading.Lock()
_upload_rl_buckets: dict[str, collections.deque] = {}
_UPLOAD_RL_MAX = 10     # max uploads
_UPLOAD_RL_WINDOW = 60  # seconds


def _check_upload_rate_limit(user_id: str) -> None:
    now = time.time()
    with _upload_rl_lock:
        bucket = _upload_rl_buckets.setdefault(user_id, collections.deque())
        # Drop timestamps outside the window
        while bucket and bucket[0] <= now - _UPLOAD_RL_WINDOW:
            bucket.popleft()
        if len(bucket) >= _UPLOAD_RL_MAX:
            raise HTTPException(
                status_code=429,
                detail=f"Слишком много загрузок. Лимит: {_UPLOAD_RL_MAX} файлов в минуту.",
            )
        bucket.append(now)


# ── Route-level helper (raises HTTPException вместо None) ────────────────────

def _get_collection(name: str):
    if name not in COLLECTIONS:
        raise HTTPException(status_code=400, detail=f"Неизвестная коллекция: {name}")
    coll = _get_collection_shared(name)
    if coll is None:
        raise HTTPException(status_code=503, detail="ChromaDB недоступна")
    return coll


# ── Фоновый подсчёт документов ──────────────────────────────────────────────

def _load_doc_counts() -> None:
    """Поднять кеш с диска при старте (best-effort)."""
    global _doc_counts, _doc_counts_at
    try:
        raw = json.loads(_DOC_COUNT_FILE.read_text(encoding="utf-8"))
        _doc_counts = {str(k): int(v) for k, v in (raw.get("counts") or {}).items()}
        _doc_counts_at = float(raw.get("computed_at") or 0.0)
    except (OSError, ValueError, TypeError) as e:
        logger.debug(f"RAG doc-count кеш не поднят с диска: {e}")


def _save_doc_counts(counts: Dict[str, int], computed_at: float) -> None:
    try:
        _DOC_COUNT_FILE.parent.mkdir(parents=True, exist_ok=True)
        _DOC_COUNT_FILE.write_text(
            json.dumps({"counts": counts, "computed_at": computed_at}, ensure_ascii=False),
            encoding="utf-8",
        )
    except OSError as e:
        logger.warning(f"Не удалось сохранить RAG doc-count кеш: {e}")


def _count_docs(coll) -> int:
    """Уникальные doc_id в коллекции. Дорого — только для фонового потока."""
    total = coll.count()
    doc_ids: set = set()
    offset = 0
    while offset < total:
        batch = coll.get(include=["metadatas"], limit=_DOC_COUNT_PAGE, offset=offset)
        metas = batch.get("metadatas") or []
        if not metas:
            break
        doc_ids.update(m.get("doc_id") for m in metas if m.get("doc_id"))
        offset += len(metas)
    return len(doc_ids)


def _refresh_doc_counts() -> None:
    global _doc_counts, _doc_counts_at, _doc_count_running
    try:
        counts: Dict[str, int] = {}
        for name in COLLECTIONS:
            coll = _get_collection_shared(name)
            if coll is None:
                continue
            counts[name] = _count_docs(coll)
        now = time.time()
        with _doc_count_lock:
            _doc_counts = counts
            _doc_counts_at = now
        _save_doc_counts(counts, now)
        logger.info(f"RAG doc-count обновлён: {counts}")
    except Exception as e:
        logger.warning(f"Фоновый пересчёт RAG doc-count не удался: {e}")
    finally:
        with _doc_count_lock:
            _doc_count_running = False


def _fresh() -> bool:
    with _doc_count_lock:
        return (time.time() - _doc_counts_at) < _DOC_COUNT_TTL


def _ensure_doc_counts_fresh() -> None:
    """Запустить фоновый пересчёт, если кеш протух. Не блокирует запрос."""
    global _doc_count_running
    with _doc_count_lock:
        if _doc_count_running:
            return
    # Файл мог появиться/обновиться уже ПОСЛЕ импорта модуля — например, кеш
    # прогрели отдельным процессом при деплое. Без этой перечитки живой сервер
    # показывал бы нули до конца фонового прохода, хотя готовые цифры уже на диске.
    if not _fresh():
        _load_doc_counts()
    if _fresh():
        return
    with _doc_count_lock:
        if _doc_count_running:
            return
        _doc_count_running = True
    threading.Thread(target=_refresh_doc_counts, name="rag-doc-counts", daemon=True).start()


_load_doc_counts()


# ── Text extraction ──────────────────────────────────────────────────────────

def _extract_text(content: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".txt":
        return content.decode("utf-8", errors="ignore")
    if ext == ".pdf":
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    if ext in (".docx", ".doc"):
        import docx
        doc = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    raise HTTPException(
        status_code=400,
        detail=f"Неподдерживаемый формат: {ext}. Используйте .txt, .pdf, .docx",
    )


def _chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> List[str]:
    sentences = re.split(r"(?<=[.!?])\s+", text)
    chunks, current, current_len = [], [], 0
    for sentence in sentences:
        slen = len(sentence)
        if current_len + slen > chunk_size and current:
            chunks.append(" ".join(current))
            overlap_text = " ".join(current)[-overlap:]
            current = [overlap_text, sentence]
            current_len = len(overlap_text) + slen
        else:
            current.append(sentence)
            current_len += slen
    if current:
        chunks.append(" ".join(current))
    return [c for c in chunks if len(c.strip()) >= 50]


# ── Schemas ──────────────────────────────────────────────────────────────────

class CollectionStat(BaseModel):
    name: str
    label: str
    chunk_count: int
    doc_count: int
    # False, пока фоновый пересчёт не отработал ни разу. Без этого флага UI рисовал
    # крупный «0» в плитке — визуально неотличимо от «данных нет», хотя чанки уже
    # посчитаны. Именно так админка выглядела пустой при живой базе знаний.
    doc_count_ready: bool = True


class RAGDocument(BaseModel):
    doc_id: str
    title: str
    collection: str
    doc_type: Optional[str] = None
    chunks: int
    uploaded_by: Optional[str] = None
    created_at: Optional[str] = None


class StatsResponse(BaseModel):
    collections: List[CollectionStat]


class DocumentsResponse(BaseModel):
    documents: List[RAGDocument]
    total: int


# ── Endpoints ────────────────────────────────────────────────────────────────

class UploadResponse(BaseModel):
    ok: bool
    doc_id: str
    title: str
    collection: str
    chunks: int


class DeleteResponse(BaseModel):
    ok: bool
    deleted_chunks: int


@router.get("/stats", response_model=StatsResponse)
async def get_stats(current_user: User = Depends(get_current_user)):
    """Статистика по всем коллекциям ChromaDB."""
    global _stats_cache
    if _stats_cache is not None and _stats_cache[1] > time.time():
        return _stats_cache[0]

    # Дорогой подсчёт doc_count — фоном; сам ответ собираем только из дешёвых
    # count(). Раньше обход метаданных всех коллекций (~1.3 млн чанков) занимал
    # больше минуты и не укладывался в 30-секундный таймаут фронта → панель
    # оставалась пустой, хотя данные в сторе были на месте.
    _ensure_doc_counts_fresh()
    with _doc_count_lock:
        known_docs = dict(_doc_counts)

    result = []
    for name in COLLECTIONS:
        chunk_count = 0
        try:
            chunk_count = _get_collection(name).count()
        except HTTPException as e:
            logger.warning(f"Коллекция {name} недоступна: {e.detail}")
        except Exception as e:
            logger.warning(f"Не удалось получить статистику для {name}: {e}")
        result.append(CollectionStat(
            name=name,
            label=COLLECTION_LABELS.get(name, name),
            chunk_count=chunk_count,
            doc_count=known_docs.get(name, 0),
            doc_count_ready=name in known_docs,
        ))
    response = StatsResponse(collections=result)
    _stats_cache = (response, time.time() + _STATS_TTL)
    return response


@router.get("/documents", response_model=DocumentsResponse)
async def list_documents(
    collection: str = Query("knowledge", description="Коллекция ChromaDB"),
    limit: int = Query(50, ge=1, le=200, description="Максимум документов"),
    offset: int = Query(0, ge=0, description="Смещение (для пагинации)"),
    current_user: User = Depends(get_current_user),
):
    """Список документов в коллекции (сгруппированных по doc_id)."""
    coll = _get_collection(collection)

    # Раньше здесь был coll.get(include=["metadatas"]) БЕЗ лимита: для laws это
    # 375 тыс. метадатных строк разом — запрос не укладывался в таймаут фронта и
    # съедал память. Читаем страницами и ограничиваем глубину: пользовательские
    # коллекции (knowledge/templates) малы и вычитываются целиком, а массовые
    # laws/case_law не подвешивают панель.
    is_admin = current_user.role in ("admin", "senior_lawyer")
    total_chunks = coll.count()
    docs: Dict[str, Dict[str, Any]] = {}
    scanned = 0
    truncated = False
    while scanned < total_chunks:
        if scanned >= _DOC_SCAN_MAX_CHUNKS:
            truncated = True
            break
        batch = coll.get(include=["metadatas"], limit=_DOC_COUNT_PAGE, offset=scanned)
        metas = batch.get("metadatas") or []
        if not metas:
            break
        for meta in metas:
            doc_id = meta.get("doc_id")
            if not doc_id:
                continue
            # IDOR fix: non-admin users see only their own documents
            if not is_admin and meta.get("uploaded_by") != str(current_user.id):
                continue
            if doc_id not in docs:
                docs[doc_id] = {
                    "doc_id": doc_id,
                    "title": meta.get("title", doc_id),
                    "collection": collection,
                    "doc_type": meta.get("doc_type"),
                    "uploaded_by": meta.get("uploaded_by"),
                    "created_at": meta.get("created_at"),
                    "chunks": 0,
                }
            docs[doc_id]["chunks"] += 1
        scanned += len(metas)

    all_documents = [RAGDocument(**d) for d in docs.values()]
    all_documents.sort(key=lambda d: d.created_at or "", reverse=True)
    paginated = all_documents[offset : offset + limit]

    # total — ровно то, что реально доступно для листания. Подставлять сюда
    # полный doc_count из кеша нельзя: при обрыве обхода по лимиту пагинация
    # уводила бы на заведомо пустые страницы. Полные цифры показывает /stats.
    if truncated:
        logger.info(
            f"Коллекция {collection}: обход прерван на {scanned} чанках "
            f"(лимит {_DOC_SCAN_MAX_CHUNKS}), показаны {len(all_documents)} док."
        )
    return DocumentsResponse(documents=paginated, total=len(all_documents))


@router.post("/documents", status_code=status.HTTP_201_CREATED, response_model=UploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    collection: str = Query("knowledge"),
    doc_type: Optional[str] = Query(None, description="Тип документа"),
    current_user: User = Depends(get_current_user),
):
    """
    Загрузить документ в ChromaDB.
    Поддерживаемые форматы: .txt, .pdf, .docx
    """
    _check_upload_rate_limit(str(current_user.id))

    # Ограниченное чтение: раньше `await file.read()` тянул весь файл в память без
    # лимита (memory-DoS — гигабайтный файл ронял воркер). Читаем максимум
    # RAG_MAX_FILE_SIZE+1 байт и отбиваем превышение (в отличие от contracts/upload
    # тут лимита не было вовсе).
    content = await file.read(RAG_MAX_FILE_SIZE + 1)
    if len(content) > RAG_MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"Файл слишком большой (>{RAG_MAX_FILE_SIZE // (1024 * 1024)} MB)",
        )
    if not content:
        raise HTTPException(status_code=400, detail="Файл пустой")

    text = _extract_text(content, file.filename or "document.txt")
    if not text.strip():
        raise HTTPException(status_code=400, detail="Не удалось извлечь текст из файла")

    doc_id = hashlib.md5(f"{file.filename}_{current_user.id}".encode()).hexdigest()
    chunks = _chunk_text(text)

    if not chunks:
        raise HTTPException(status_code=400, detail="Документ слишком короткий для индексации")

    now = datetime.now(timezone.utc).isoformat()
    title = Path(file.filename or "document").stem
    coll = _get_collection(collection)

    existing = coll.get(where={"doc_id": doc_id})
    if existing["ids"]:
        coll.delete(ids=existing["ids"])

    ids = [f"{doc_id}_chunk_{i}" for i in range(len(chunks))]
    metadatas = [
        {
            "doc_id": doc_id,
            "title": title,
            "doc_type": doc_type or "document",
            "chunk_id": i,
            "total_chunks": len(chunks),
            "uploaded_by": str(current_user.id),
            "created_at": now,
            "filename": file.filename or "",
        }
        for i in range(len(chunks))
    ]

    coll.add(ids=ids, documents=chunks, metadatas=metadatas)
    logger.info(
        f"RAG: загружен '{title}' ({len(chunks)} чанков) → '{collection}' пользователем {current_user.id}"
    )
    global _stats_cache
    _stats_cache = None  # invalidate stats cache after upload

    return UploadResponse(ok=True, doc_id=doc_id, title=title, collection=collection, chunks=len(chunks))


@router.delete("/documents/{doc_id}", response_model=DeleteResponse)
async def delete_document(
    doc_id: str,
    collection: str = Query("knowledge"),
    current_user: User = Depends(get_current_user),
):
    """Удалить документ из ChromaDB по doc_id."""
    coll = _get_collection(collection)
    existing = coll.get(where={"doc_id": doc_id}, include=["metadatas"])

    if not existing["ids"]:
        raise HTTPException(status_code=404, detail="Документ не найден в коллекции")

    # Verify ownership — only the uploader or an admin can delete (IDOR prevention)
    if current_user.role != "admin":
        meta = existing["metadatas"][0] if existing["metadatas"] else {}
        if meta.get("uploaded_by") != str(current_user.id):
            raise HTTPException(status_code=403, detail="Нет доступа к этому документу")

    coll.delete(ids=existing["ids"])
    logger.info(f"RAG: удалён документ {doc_id} из коллекции '{collection}' пользователем {current_user.id}")
    global _stats_cache
    _stats_cache = None  # invalidate stats cache after delete
    return DeleteResponse(ok=True, deleted_chunks=len(existing["ids"]))
