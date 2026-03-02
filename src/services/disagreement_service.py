# -*- coding: utf-8 -*-
"""
Сервис генерации протоколов разногласий через LLM
Принимает результат анализа + выбранные пункты → LLM генерирует возражения → DOCX
"""
import os
import json
from typing import Dict, Any, Optional, List
from datetime import datetime
from dataclasses import dataclass, field
from loguru import logger


@dataclass
class RiskItem:
    """Риск из результата анализа."""
    section_number: str = ""
    section_title: str = ""
    original_text: str = ""
    risk_type: str = ""
    severity: str = "medium"
    description: str = ""
    consequences: str = ""
    relevant_laws: List[str] = field(default_factory=list)


@dataclass
class Objection:
    """Сгенерированное возражение к пункту договора."""
    section_number: str = ""
    section_title: str = ""
    original_text: str = ""
    issue_description: str = ""
    legal_basis: str = ""
    risk_explanation: str = ""
    proposed_formulation: str = ""
    reasoning: str = ""
    priority: str = "medium"


@dataclass
class DisagreementProtocol:
    """Полный протокол разногласий."""
    success: bool = True
    contract_name: str = ""
    contract_type: str = ""
    date: str = ""
    objections: List[Objection] = field(default_factory=list)
    docx_path: str = ""
    json_path: str = ""
    error: str = ""
    tokens_used: int = 0
    generation_time: float = 0.0


class DisagreementService:
    """
    Генерация протоколов разногласий на основе анализа договора.

    Использование:
        service = DisagreementService()
        risks = [RiskItem(section_number="3.1", description="Нет срока оплаты", severity="high")]
        protocol = service.generate(risks, contract_name="Договор поставки №123")
    """

    SYSTEM_PROMPT = """Ты — опытный юрист по договорному праву РФ. Твоя задача — составить
формальные юридические возражения к проблемным пунктам договора.

Для КАЖДОГО пункта составь:
1. **Замечание** — описание проблемы с юридической точки зрения
2. **Правовое обоснование** — ссылки на конкретные статьи ГК РФ, федеральные законы
3. **Объяснение рисков** — какие последствия возникнут, если пункт останется без изменений
4. **Предлагаемая формулировка** — конкретный текст, которым следует заменить проблемный пункт
5. **Обоснование предложения** — почему предложенная формулировка лучше

ТОН: Деловой, нейтральный, аргументированный.

ФОРМАТ ОТВЕТА: JSON-массив объектов:
[
  {
    "section_number": "3.1",
    "issue_description": "...",
    "legal_basis": "...",
    "risk_explanation": "...",
    "proposed_formulation": "...",
    "reasoning": "..."
  }
]

Отвечай ТОЛЬКО JSON-массивом, без дополнительного текста."""

    def __init__(self, provider: Optional[str] = None, exports_dir: str = "data/exports/disagreements"):
        self.provider = provider
        self.exports_dir = exports_dir
        os.makedirs(self.exports_dir, exist_ok=True)

    def generate(
        self,
        risks: List[RiskItem],
        contract_name: str = "",
        contract_type: str = "",
        selected_indices: Optional[List[int]] = None,
    ) -> DisagreementProtocol:
        """Генерирует протокол разногласий по выбранным рискам."""
        import time
        start_time = time.time()

        try:
            # Фильтруем по выбранным индексам
            if selected_indices is not None:
                selected_risks = [risks[i] for i in selected_indices if i < len(risks)]
            else:
                selected_risks = risks

            if not selected_risks:
                return DisagreementProtocol(success=False, error="Не выбрано ни одного пункта для возражений")

            # Формируем промпт
            prompt = self._build_prompt(selected_risks, contract_name, contract_type)
            logger.info(f"Генерация протокола разногласий: {len(selected_risks)} пунктов | {contract_name}")

            # Вызов LLM
            from src.services.llm_gateway import LLMGateway
            gateway = LLMGateway(provider=self.provider)

            response = gateway.call(
                prompt=prompt,
                system_prompt=self.SYSTEM_PROMPT,
                response_format="json",
                temperature=0.2,
                max_tokens=6000,
            )

            # Парсинг ответа
            objections = self._parse_response(response, selected_risks)

            if not objections:
                return DisagreementProtocol(success=False, error="LLM не вернула возражений")

            # Формируем протокол
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            protocol = DisagreementProtocol(
                success=True,
                contract_name=contract_name,
                contract_type=contract_type,
                date=datetime.now().strftime("%d.%m.%Y"),
                objections=objections,
                tokens_used=gateway.total_input_tokens + gateway.total_output_tokens,
                generation_time=time.time() - start_time,
            )

            # Сохраняем DOCX
            docx_filename = f"protocol_{timestamp}.docx"
            docx_path = os.path.join(self.exports_dir, docx_filename)
            self._save_to_docx(protocol, docx_path)
            protocol.docx_path = docx_path

            # Сохраняем JSON
            json_filename = f"protocol_{timestamp}.json"
            json_path = os.path.join(self.exports_dir, json_filename)
            self._save_to_json(protocol, json_path)
            protocol.json_path = json_path

            elapsed = time.time() - start_time
            logger.info(f"Протокол сгенерирован за {elapsed:.1f}с: {len(objections)} возражений")

            return protocol

        except Exception as e:
            logger.error(f"Ошибка генерации протокола: {e}")
            return DisagreementProtocol(success=False, error=str(e))

    def _build_prompt(self, risks: List[RiskItem], contract_name: str, contract_type: str) -> str:
        """Формирует промпт из списка рисков."""
        parts = []
        if contract_name:
            parts.append(f"ДОГОВОР: {contract_name}")
        if contract_type:
            parts.append(f"ТИП: {contract_type}")
        parts.append(f"\nПРОБЛЕМНЫЕ ПУНКТЫ ДЛЯ ВОЗРАЖЕНИЙ ({len(risks)} шт.):\n")

        for i, risk in enumerate(risks, 1):
            parts.append(f"--- Пункт {i} ---")
            if risk.section_number:
                parts.append(f"Номер раздела: {risk.section_number}")
            if risk.section_title:
                parts.append(f"Название: {risk.section_title}")
            if risk.original_text:
                parts.append(f"Текст пункта: {risk.original_text[:500]}")
            if risk.description:
                parts.append(f"Выявленная проблема: {risk.description}")
            if risk.severity:
                parts.append(f"Уровень риска: {risk.severity}")
            if risk.consequences:
                parts.append(f"Возможные последствия: {risk.consequences}")
            if risk.relevant_laws:
                parts.append(f"Связанные нормы: {', '.join(risk.relevant_laws)}")
            parts.append("")

        parts.append("Составь возражения ко ВСЕМ указанным пунктам.")
        return "\n".join(parts)

    def _parse_response(self, response: Any, risks: List[RiskItem]) -> List[Objection]:
        """Парсит ответ LLM в список возражений."""
        # Преобразуем в список словарей
        if isinstance(response, str):
            try:
                data = json.loads(response)
            except json.JSONDecodeError:
                logger.error("Не удалось распарсить JSON ответ LLM")
                return []
        elif isinstance(response, list):
            data = response
        elif isinstance(response, dict):
            data = response.get("objections", [response])
        else:
            return []

        if not isinstance(data, list):
            data = [data]

        objections = []
        for i, item in enumerate(data):
            if not isinstance(item, dict):
                continue

            # Найти соответствующий риск для дополнения данных
            risk = risks[i] if i < len(risks) else RiskItem()

            obj = Objection(
                section_number=item.get("section_number", risk.section_number),
                section_title=item.get("section_title", risk.section_title),
                original_text=risk.original_text,
                issue_description=item.get("issue_description", ""),
                legal_basis=item.get("legal_basis", ""),
                risk_explanation=item.get("risk_explanation", ""),
                proposed_formulation=item.get("proposed_formulation", ""),
                reasoning=item.get("reasoning", ""),
                priority=risk.severity,
            )
            objections.append(obj)

        return objections

    def _save_to_docx(self, protocol: DisagreementProtocol, file_path: str):
        """Сохраняет протокол разногласий в DOCX (таблица)."""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # Шрифт
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

        # Заголовок
        title = doc.add_heading("ПРОТОКОЛ РАЗНОГЛАСИЙ", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Метаданные
        if protocol.contract_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"к {protocol.contract_name}")
            run.font.size = Pt(12)

        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.add_run(f"Дата: {protocol.date}")

        doc.add_paragraph()

        # Сводка
        p_summary = doc.add_paragraph()
        run = p_summary.add_run(f"Всего замечаний: {len(protocol.objections)}")
        run.bold = True

        priority_counts = {}
        for obj in protocol.objections:
            priority_counts[obj.priority] = priority_counts.get(obj.priority, 0) + 1

        priority_labels = {"critical": "критических", "high": "высоких", "medium": "средних", "low": "низких"}
        counts_str = ", ".join(
            f"{count} {priority_labels.get(p, p)}"
            for p, count in sorted(priority_counts.items(), key=lambda x: ["critical", "high", "medium", "low"].index(x[0]) if x[0] in ["critical", "high", "medium", "low"] else 99)
        )
        if counts_str:
            doc.add_paragraph(f"Приоритет: {counts_str}")

        doc.add_paragraph()

        # Таблица возражений
        cols_count = 5
        table = doc.add_table(rows=1, cols=cols_count)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заголовки таблицы
        headers = ["№ п/п", "Пункт договора", "Замечание", "Правовое обоснование", "Предлагаемая формулировка"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Строки с возражениями
        priority_colors = {
            "critical": RGBColor(220, 20, 60),
            "high": RGBColor(255, 140, 0),
            "medium": RGBColor(180, 150, 0),
            "low": RGBColor(60, 179, 113),
        }

        for i, obj in enumerate(protocol.objections, 1):
            row = table.add_row()

            # Номер
            row.cells[0].text = str(i)

            # Пункт договора
            section_text = obj.section_number
            if obj.section_title:
                section_text += f"\n{obj.section_title}"
            row.cells[1].text = section_text

            # Замечание + риски
            issue_text = obj.issue_description
            if obj.risk_explanation:
                issue_text += f"\n\nРиски: {obj.risk_explanation}"
            row.cells[2].text = issue_text

            # Правовое обоснование
            legal_text = obj.legal_basis
            if obj.reasoning:
                legal_text += f"\n\n{obj.reasoning}"
            row.cells[3].text = legal_text

            # Предлагаемая формулировка
            row.cells[4].text = obj.proposed_formulation

            # Цвет приоритета для первого столбца
            color = priority_colors.get(obj.priority)
            if color:
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = color

        # Подпись
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("С уважением,")
        doc.add_paragraph("[Подпись]")
        doc.add_paragraph("[Должность]")
        doc.add_paragraph("[Дата]")

        doc.save(file_path)
        logger.info(f"Протокол DOCX сохранён: {file_path}")

    def _save_to_json(self, protocol: DisagreementProtocol, file_path: str):
        """Сохраняет протокол в JSON."""
        data = {
            "contract_name": protocol.contract_name,
            "contract_type": protocol.contract_type,
            "date": protocol.date,
            "total_objections": len(protocol.objections),
            "generation_time_sec": round(protocol.generation_time, 2),
            "objections": [
                {
                    "section_number": obj.section_number,
                    "section_title": obj.section_title,
                    "original_text": obj.original_text,
                    "issue_description": obj.issue_description,
                    "legal_basis": obj.legal_basis,
                    "risk_explanation": obj.risk_explanation,
                    "proposed_formulation": obj.proposed_formulation,
                    "reasoning": obj.reasoning,
                    "priority": obj.priority,
                }
                for obj in protocol.objections
            ],
        }

        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        logger.info(f"Протокол JSON сохранён: {file_path}")


__all__ = ["DisagreementService", "RiskItem", "Objection", "DisagreementProtocol"]
