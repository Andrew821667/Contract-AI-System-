"""
Document Parser Service - конвертация документов в XML
Поддержка: DOCX, PDF → XML
"""
import os
import re
import uuid
import hashlib
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
from lxml import etree
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
import pypdf
import pdfplumber
from loguru import logger

from ..utils.xml_security import parse_xml_safely, XMLSecurityError


def _get_redis():
    """Get Redis connection (lazy, returns None if unavailable)"""
    try:
        import redis as redis_lib
        from config.settings import settings
        return redis_lib.from_url(settings.redis_url, decode_responses=True)
    except Exception:
        return None


class DocumentParser:
    """Парсер документов в XML формат"""

    XML_CACHE_TTL = 3600  # 1 hour

    def __init__(self):
        self.supported_formats = ['.docx', '.pdf', '.txt']
        self._ocr_service = None
        self._redis = _get_redis()

    def _get_ocr_service(self):
        """Lazy-init OCR service (only when needed)"""
        if self._ocr_service is None:
            try:
                from .ocr_service import OCRService
                self._ocr_service = OCRService()
            except Exception as e:
                logger.warning(f"OCR service not available: {e}")
                self._ocr_service = False  # Sentinel: tried but failed
        return self._ocr_service if self._ocr_service is not False else None

    def _file_cache_key(self, file_path: str) -> Optional[str]:
        """Compute SHA256 of file contents for cache key"""
        try:
            h = hashlib.sha256()
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(8192), b''):
                    h.update(chunk)
            return f"xml_cache:{h.hexdigest()}"
        except Exception:
            return None

    def parse(self, file_path: str) -> str:
        """
        Универсальный парсинг файла в XML

        Args:
            file_path: Путь к файлу

        Returns:
            XML строка договора

        Raises:
            ValueError: Если формат не поддерживается
            FileNotFoundError: Если файл не найден
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = Path(file_path).suffix.lower()

        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {ext}. Supported: {self.supported_formats}")

        # Check Redis cache
        cache_key = self._file_cache_key(file_path)
        if cache_key and self._redis:
            try:
                cached = self._redis.get(cache_key)
                if cached:
                    logger.info(f"XML cache HIT for {file_path}")
                    return cached
            except Exception as e:
                logger.debug(f"Redis cache read failed: {e}")

        logger.info(f"Parsing document: {file_path} (format: {ext})")

        if ext == '.docx':
            result = self.parse_docx(file_path)
        elif ext == '.pdf':
            result = self.parse_pdf(file_path)
        elif ext == '.txt':
            result = self.parse_txt(file_path)
        else:
            result = None

        # Store in Redis cache
        if result and cache_key and self._redis:
            try:
                self._redis.setex(cache_key, self.XML_CACHE_TTL, result)
                logger.info(f"XML cache SAVE for {file_path}")
            except Exception as e:
                logger.debug(f"Redis cache write failed: {e}")

        return result

    def parse_docx(self, docx_path: str) -> str:
        """
        Парсинг DOCX в XML

        Args:
            docx_path: Путь к DOCX файлу

        Returns:
            XML строка
        """
        logger.info(f"Parsing DOCX: {docx_path}")

        doc = Document(docx_path)

        # Создаём корневой элемент
        root = etree.Element("contract")

        # Метаданные
        metadata = self._extract_metadata_docx(doc, docx_path)
        metadata_elem = etree.SubElement(root, "metadata")
        for key, value in metadata.items():
            elem = etree.SubElement(metadata_elem, key)
            elem.text = str(value)

        # Извлекаем структуру документа
        sections = self._extract_sections_docx(doc)

        # Пытаемся извлечь стороны договора
        parties = self._extract_parties_from_text(sections)
        if parties:
            parties_elem = etree.SubElement(root, "parties")
            for party in parties:
                party_elem = etree.SubElement(parties_elem, "party")
                party_elem.set("role", party.get("role", "unknown"))
                for key, value in party.items():
                    if key != "role":
                        elem = etree.SubElement(party_elem, key)
                        elem.text = str(value)

        # Пытаемся извлечь финансовые условия
        terms = self._extract_terms_from_text(sections)
        if terms:
            terms_elem = etree.SubElement(root, "terms")

            # Финансовые условия
            if terms.get("financial"):
                financial_elem = etree.SubElement(terms_elem, "financial")
                for key, value in terms["financial"].items():
                    elem = etree.SubElement(financial_elem, key)
                    elem.text = str(value)

            # Даты
            if terms.get("dates"):
                dates_elem = etree.SubElement(terms_elem, "dates")
                for key, value in terms["dates"].items():
                    elem = etree.SubElement(dates_elem, key)
                    elem.text = str(value)

        # Разделы (clauses)
        clauses_elem = etree.SubElement(root, "clauses")
        for idx, section in enumerate(sections, 1):
            clause_elem = etree.SubElement(clauses_elem, "clause")
            clause_elem.set("id", str(idx))
            clause_elem.set("type", section.get("type", "general"))

            title_elem = etree.SubElement(clause_elem, "title")
            title_elem.text = section.get("title", f"Раздел {idx}")

            content_elem = etree.SubElement(clause_elem, "content")
            for para in section.get("paragraphs", []):
                para_elem = etree.SubElement(content_elem, "paragraph")
                para_elem.text = para

        # Таблицы (если есть)
        tables = self._extract_tables_docx(doc)
        if tables:
            tables_elem = etree.SubElement(root, "tables")
            for idx, table_data in enumerate(tables, 1):
                table_elem = etree.SubElement(tables_elem, "table")
                table_elem.set("id", str(idx))

                for row_data in table_data:
                    row_elem = etree.SubElement(table_elem, "row")
                    for cell_text in row_data:
                        cell_elem = etree.SubElement(row_elem, "cell")
                        cell_elem.text = cell_text

        # Конвертируем в строку с форматированием
        xml_str = etree.tostring(
            root,
            encoding='unicode',
            pretty_print=True
        )
        # Добавляем XML declaration вручную
        xml_str = '<?xml version="1.0" encoding="UTF-8"?>\n' + xml_str

        logger.info(f"DOCX parsed successfully: {len(sections)} sections")
        return xml_str

    def parse_pdf(self, pdf_path: str) -> str:
        """
        Парсинг PDF в XML

        Args:
            pdf_path: Путь к PDF файлу

        Returns:
            XML строка
        """
        logger.info(f"Parsing PDF: {pdf_path}")

        # Извлекаем текст с помощью pdfplumber (лучше сохраняет структуру)
        text_content = []

        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying pypdf: {e}")
            # Fallback на pypdf
            with open(pdf_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)

        full_text = "\n\n".join(text_content)

        # If very little text extracted, try OCR (likely a scanned PDF)
        if len(full_text.strip()) < 100:
            ocr = self._get_ocr_service()
            if ocr:
                try:
                    is_scanned, reason = ocr.detect_if_scanned(pdf_path)
                    if is_scanned:
                        logger.info(f"Scanned PDF detected ({reason}), running OCR...")
                        ocr_text = ocr.extract_text_from_pdf(pdf_path)
                        if len(ocr_text.strip()) > len(full_text.strip()):
                            full_text = ocr_text
                            logger.info(f"OCR extracted {len(ocr_text)} chars from scanned PDF")
                except Exception as e:
                    logger.warning(f"OCR fallback failed: {e}")

        # Создаём структуру из текста
        sections = self._extract_sections_from_text(full_text)

        # Создаём XML
        root = etree.Element("contract")

        # Метаданные
        metadata = self._extract_metadata_pdf(pdf_path, full_text)
        metadata_elem = etree.SubElement(root, "metadata")
        for key, value in metadata.items():
            elem = etree.SubElement(metadata_elem, key)
            elem.text = str(value)

        # Стороны
        parties = self._extract_parties_from_text(sections)
        if parties:
            parties_elem = etree.SubElement(root, "parties")
            for party in parties:
                party_elem = etree.SubElement(parties_elem, "party")
                party_elem.set("role", party.get("role", "unknown"))
                for key, value in party.items():
                    if key != "role":
                        elem = etree.SubElement(party_elem, key)
                        elem.text = str(value)

        # Условия
        terms = self._extract_terms_from_text(sections)
        if terms:
            terms_elem = etree.SubElement(root, "terms")

            if terms.get("financial"):
                financial_elem = etree.SubElement(terms_elem, "financial")
                for key, value in terms["financial"].items():
                    elem = etree.SubElement(financial_elem, key)
                    elem.text = str(value)

            if terms.get("dates"):
                dates_elem = etree.SubElement(terms_elem, "dates")
                for key, value in terms["dates"].items():
                    elem = etree.SubElement(dates_elem, key)
                    elem.text = str(value)

        # Разделы
        clauses_elem = etree.SubElement(root, "clauses")
        for idx, section in enumerate(sections, 1):
            clause_elem = etree.SubElement(clauses_elem, "clause")
            clause_elem.set("id", str(idx))
            clause_elem.set("type", section.get("type", "general"))

            title_elem = etree.SubElement(clause_elem, "title")
            title_elem.text = section.get("title", f"Раздел {idx}")

            content_elem = etree.SubElement(clause_elem, "content")
            for para in section.get("paragraphs", []):
                para_elem = etree.SubElement(content_elem, "paragraph")
                para_elem.text = para

        xml_str = etree.tostring(
            root,
            encoding='unicode',
            pretty_print=True
        )
        # Добавляем XML declaration вручную
        xml_str = '<?xml version="1.0" encoding="UTF-8"?>\n' + xml_str

        logger.info(f"PDF parsed successfully: {len(sections)} sections")
        return xml_str

    def parse_txt(self, txt_path: str) -> str:
        """
        Парсинг TXT в XML

        Args:
            txt_path: Путь к TXT файлу

        Returns:
            XML строка
        """
        logger.info(f"Parsing TXT: {txt_path}")

        with open(txt_path, 'r', encoding='utf-8', errors='replace') as f:
            full_text = f.read()

        sections = self._extract_sections_from_text(full_text)

        root = etree.Element("contract")

        metadata = {
            "contract_id": str(uuid.uuid4()),
            "file_name": Path(txt_path).name,
            "creation_date": datetime.now().isoformat(),
            "version": "1.0",
            "format": "txt",
            "text_length": str(len(full_text))
        }
        metadata_elem = etree.SubElement(root, "metadata")
        for key, value in metadata.items():
            elem = etree.SubElement(metadata_elem, key)
            elem.text = str(value)

        parties = self._extract_parties_from_text(sections)
        if parties:
            parties_elem = etree.SubElement(root, "parties")
            for party in parties:
                party_elem = etree.SubElement(parties_elem, "party")
                party_elem.set("role", party.get("role", "unknown"))
                for key, value in party.items():
                    if key != "role":
                        elem = etree.SubElement(party_elem, key)
                        elem.text = str(value)

        terms = self._extract_terms_from_text(sections)
        if terms:
            terms_elem = etree.SubElement(root, "terms")
            if terms.get("financial"):
                financial_elem = etree.SubElement(terms_elem, "financial")
                for key, value in terms["financial"].items():
                    elem = etree.SubElement(financial_elem, key)
                    elem.text = str(value)
            if terms.get("dates"):
                dates_elem = etree.SubElement(terms_elem, "dates")
                for key, value in terms["dates"].items():
                    elem = etree.SubElement(dates_elem, key)
                    elem.text = str(value)

        clauses_elem = etree.SubElement(root, "clauses")
        for idx, section in enumerate(sections, 1):
            clause_elem = etree.SubElement(clauses_elem, "clause")
            clause_elem.set("id", str(idx))
            clause_elem.set("type", section.get("type", "general"))

            title_elem = etree.SubElement(clause_elem, "title")
            title_elem.text = section.get("title", f"Раздел {idx}")

            content_elem = etree.SubElement(clause_elem, "content")
            for para in section.get("paragraphs", []):
                para_elem = etree.SubElement(content_elem, "paragraph")
                para_elem.text = para

        xml_str = etree.tostring(root, encoding='unicode', pretty_print=True)
        xml_str = '<?xml version="1.0" encoding="UTF-8"?>\n' + xml_str

        logger.info(f"TXT parsed successfully: {len(sections)} sections")
        return xml_str

    def extract_tracked_changes(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        Извлечение tracked changes из DOCX

        Args:
            docx_path: Путь к DOCX файлу

        Returns:
            Список изменений
        """
        logger.info(f"Extracting tracked changes from: {docx_path}")

        doc = Document(docx_path)
        changes = []

        # Проходим по всем параграфам
        for para_idx, paragraph in enumerate(doc.paragraphs):
            # Получаем XML элемент параграфа
            p_element = paragraph._element

            # Ищем insertions (w:ins)
            for ins in p_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ins'):
                change = {
                    "type": "insert",
                    "paragraph_id": para_idx,
                    "author": ins.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown'),
                    "date": ins.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', ''),
                    "text": self._get_text_from_element(ins)
                }
                changes.append(change)

            # Ищем deletions (w:del)
            for delete in p_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}del'):
                change = {
                    "type": "delete",
                    "paragraph_id": para_idx,
                    "author": delete.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown'),
                    "date": delete.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', ''),
                    "text": self._get_text_from_element(delete)
                }
                changes.append(change)

        logger.info(f"Found {len(changes)} tracked changes")
        return changes

    def _extract_metadata_docx(self, doc: Document, file_path: str) -> Dict[str, str]:
        """Извлечение метаданных из DOCX"""
        return {
            "contract_id": str(uuid.uuid4()),
            "file_name": Path(file_path).name,
            "creation_date": datetime.now().isoformat(),
            "version": "1.0",
            "format": "docx",
            "title": doc.core_properties.title or "Без названия",
            "author": doc.core_properties.author or "Неизвестно"
        }

    def _extract_metadata_pdf(self, pdf_path: str, text: str) -> Dict[str, str]:
        """Извлечение метаданных из PDF"""
        return {
            "contract_id": str(uuid.uuid4()),
            "file_name": Path(pdf_path).name,
            "creation_date": datetime.now().isoformat(),
            "version": "1.0",
            "format": "pdf",
            "text_length": str(len(text))
        }

    def _extract_sections_docx(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Извлечение разделов из DOCX
        Определяет структуру по заголовкам (Heading 1, 2, 3)
        """
        sections = []
        current_section = None

        for para in doc.paragraphs:
            # Проверяем, является ли параграф заголовком
            if para.style.name.startswith('Heading'):
                # Сохраняем предыдущую секцию
                if current_section:
                    sections.append(current_section)

                # Начинаем новую секцию
                current_section = {
                    "title": para.text.strip(),
                    "type": self._classify_section_type(para.text),
                    "paragraphs": []
                }
            else:
                # Обычный параграф
                text = para.text.strip()
                if text:
                    if current_section is None:
                        # Если ещё нет секции, создаём дефолтную
                        current_section = {
                            "title": "Преамбула",
                            "type": "preamble",
                            "paragraphs": []
                        }
                    current_section["paragraphs"].append(text)

        # Добавляем последнюю секцию
        if current_section:
            sections.append(current_section)

        return sections

    def _extract_sections_from_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Извлечение разделов из plain text (для PDF)
        Ищет паттерны типа "1.", "1.1", "Раздел 1" и т.д.
        """
        sections = []

        # Разбиваем по паттернам заголовков
        # Паттерн: цифра с точкой в начале строки
        section_pattern = r'^(\d+\.?\s+[^\n]*)'

        lines = text.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Проверяем, является ли строка заголовком
            if re.match(section_pattern, line):
                # Сохраняем предыдущую секцию
                if current_section:
                    sections.append(current_section)

                # Новая секция
                current_section = {
                    "title": line,
                    "type": self._classify_section_type(line),
                    "paragraphs": []
                }
            else:
                # Обычная строка
                if current_section is None:
                    current_section = {
                        "title": "Преамбула",
                        "type": "preamble",
                        "paragraphs": []
                    }
                if len(line) > 10:  # Игнорируем очень короткие строки
                    current_section["paragraphs"].append(line)

        # Последняя секция
        if current_section:
            sections.append(current_section)

        return sections

    def _extract_tables_docx(self, doc: Document) -> List[List[List[str]]]:
        """Извлечение таблиц из DOCX"""
        tables = []

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append(table_data)

        return tables

    def _extract_parties_from_text(self, sections: List[Dict]) -> List[Dict[str, str]]:
        """
        Попытка извлечь стороны договора из текста
        Ищет паттерны типа "ООО", "АО", и т.д.
        """
        parties = []

        # Ищем в первых разделах упоминания организаций
        text_to_search = ""
        for section in sections[:3]:  # Проверяем первые 3 раздела
            text_to_search += " ".join(section.get("paragraphs", []))

        # Паттерн для организаций (ООО, АО, ПАО, ЗАО, ФГУП, МУП, НКО, ИП и др.)
        org_pattern = r'(ООО|ОАО|ПАО|АО|ЗАО|ФГУП|МУП|ГУП|НКО|АНО|ИП)\s+"([^"]+)"'
        matches = re.findall(org_pattern, text_to_search)

        for idx, (org_type, org_name) in enumerate(matches[:2]):  # Максимум 2 стороны
            role = "supplier" if idx == 0 else "buyer"
            parties.append({
                "role": role,
                "name": f'{org_type} "{org_name}"'
            })

        # Ищем ИНН
        inn_pattern = r'ИНН[:\s]+(\d{10,12})'
        inn_matches = re.findall(inn_pattern, text_to_search)

        for idx, inn in enumerate(inn_matches[:len(parties)]):
            if idx < len(parties):
                parties[idx]["inn"] = inn

        return parties

    def _extract_terms_from_text(self, sections: List[Dict]) -> Dict[str, Any]:
        """
        Извлечение условий из текста (финансы, даты)
        """
        terms = {
            "financial": {},
            "dates": {}
        }

        text_to_search = ""
        for section in sections:
            text_to_search += " ".join(section.get("paragraphs", []))

        # Ищем суммы (рубли)
        amount_pattern = r'(\d+(?:\s?\d+)*(?:[,\.]\d+)?)\s*(?:рублей|рубль)'
        amounts = re.findall(amount_pattern, text_to_search)
        if amounts:
            # Берём первую найденную сумму
            amount_str = amounts[0].replace(' ', '').replace(',', '.')
            try:
                terms["financial"]["total_amount"] = amount_str
                terms["financial"]["currency"] = "RUB"
            except:
                pass

        # Ищем даты (формат ДД.ММ.ГГГГ)
        date_pattern = r'(\d{2}\.\d{2}\.\d{4})'
        dates = re.findall(date_pattern, text_to_search)
        if dates:
            # Берём первые найденные даты как дату договора
            if len(dates) >= 1:
                terms["dates"]["signature_date"] = dates[0]

        return terms

    def _classify_section_type(self, title: str) -> str:
        """
        Классификация типа раздела по названию
        """
        title_lower = title.lower()

        if any(word in title_lower for word in ['предмет', 'subject']):
            return 'subject'
        elif any(word in title_lower for word in ['цена', 'расчет', 'оплата', 'payment', 'price', 'стоимость', 'вознаграждение']):
            return 'financial'
        elif any(word in title_lower for word in ['срок', 'deadline', 'term', 'период', 'действи']):
            return 'terms'
        elif any(word in title_lower for word in ['ответственность', 'liability', 'штраф', 'неустойк', 'пени']):
            return 'liability'
        elif any(word in title_lower for word in ['спор', 'dispute', 'арбитраж', 'претенз']):
            return 'dispute'
        elif any(word in title_lower for word in ['конфиденциальн', 'confidential', 'тайн']):
            return 'confidentiality'
        elif any(word in title_lower for word in ['форс-мажор', 'force majeure', 'обстоятельств']):
            return 'force_majeure'
        elif any(word in title_lower for word in ['расторж', 'прекращ', 'termination']):
            return 'termination'
        elif any(word in title_lower for word in ['гарант', 'качеств', 'warranty']):
            return 'warranty'
        else:
            return 'general'

    def _get_text_from_element(self, element) -> str:
        """Извлечение текста из XML элемента"""
        text_parts = []
        for text_node in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if text_node.text:
                text_parts.append(text_node.text)
        return ''.join(text_parts)

    def validate_xml(self, xml_str: str) -> bool:
        """
        Валидация XML

        Args:
            xml_str: XML строка

        Returns:
            True если валидна
        """
        try:
            parse_xml_safely(xml_str)
            return True
        except (etree.XMLSyntaxError, XMLSecurityError) as e:
            logger.error(f"XML validation failed: {e}")
            return False


# Экспорт
__all__ = ["DocumentParser"]
