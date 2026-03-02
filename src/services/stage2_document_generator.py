# -*- coding: utf-8 -*-
"""
Stage 2 Final Document Generator (2.4)

Provides two output variants:
- Variant A: Corrected DOCX with accepted recommendations applied to text
- Variant B: Disagreement protocol DOCX/JSON (before -> after)
"""

import io
import json
import logging
import re
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class Stage2DocumentGenerator:
    """Generate final outputs for Stage 2."""

    def generate_corrected_docx(
        self,
        base_docx_bytes: Optional[bytes],
        accepted_recommendations: List[Dict[str, Any]],
        source_file_name: str = "contract",
        raw_text: str = "",
    ) -> bytes:
        """
        Variant A:
        Build corrected DOCX — применяет рекомендации к тексту документа,
        а затем добавляет приложение со списком внесённых правок.
        """

        # 1. Получить базовый документ
        if base_docx_bytes:
            document = Document(BytesIO(base_docx_bytes))
        elif raw_text:
            # Fallback: создаём DOCX из извлечённого текста
            document = self._text_to_docx(raw_text)
        else:
            document = Document()
            document.add_heading("Договор (текст не извлечён)", level=1)

        if not accepted_recommendations:
            return self._to_bytes(document)

        # 2. Применяем рекомендации к тексту документа
        applied, not_applied = self._apply_recommendations(document, accepted_recommendations)

        # 3. Добавляем приложение со списком внесённых правок
        document.add_page_break()
        h = document.add_heading("Приложение: Перечень внесённых правок", level=1)
        document.add_paragraph(f"Документ-источник: {source_file_name}")
        document.add_paragraph(f"Дата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        document.add_paragraph(f"Всего правок: {len(accepted_recommendations)} "
                               f"(применено к тексту: {len(applied)}, "
                               f"не удалось применить автоматически: {len(not_applied)})")
        document.add_paragraph("")

        # Список применённых правок
        if applied:
            document.add_heading("Применённые правки", level=2)
            for idx, rec in enumerate(applied, 1):
                self._add_recommendation_entry(document, idx, rec, applied=True)

        # Список правок, которые не удалось применить автоматически
        if not_applied:
            document.add_heading("Правки для ручного внесения", level=2)
            p_note = document.add_paragraph()
            run = p_note.add_run(
                "Следующие правки не удалось автоматически применить к тексту. "
                "Внесите их вручную."
            )
            run.italic = True
            for idx, rec in enumerate(not_applied, 1):
                self._add_recommendation_entry(document, idx, rec, applied=False)

        return self._to_bytes(document)

    def _apply_recommendations(
        self,
        document: Document,
        recommendations: List[Dict[str, Any]],
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        Применяет рекомендации к параграфам документа.
        Возвращает (applied, not_applied).
        """
        applied: List[Dict[str, Any]] = []
        not_applied: List[Dict[str, Any]] = []

        # Собираем все параграфы с их текстом для поиска
        paragraphs = list(document.paragraphs)

        for rec in recommendations:
            action_type = (rec.get("action_type") or "modify").lower()
            original_text = (rec.get("original_text") or "").strip()
            proposed_text = (rec.get("proposed_text") or "").strip()
            section_title = (rec.get("section_title") or "").strip()

            if action_type == "modify" and original_text and proposed_text:
                # Ищем original_text в параграфах и заменяем
                found = self._find_and_replace(paragraphs, original_text, proposed_text)
                if found:
                    applied.append(rec)
                else:
                    not_applied.append(rec)

            elif action_type == "add" and proposed_text:
                # Ищем раздел по section_title и вставляем после него
                inserted = self._insert_after_section(
                    document, paragraphs, section_title, proposed_text
                )
                if inserted:
                    applied.append(rec)
                    # Обновляем список параграфов
                    paragraphs = list(document.paragraphs)
                else:
                    not_applied.append(rec)

            elif action_type == "remove" and original_text:
                # Помечаем текст зачёркиванием
                found = self._find_and_strikethrough(paragraphs, original_text)
                if found:
                    applied.append(rec)
                else:
                    not_applied.append(rec)
            else:
                # Нет достаточных данных для автоприменения
                not_applied.append(rec)

        return applied, not_applied

    def _find_and_replace(
        self,
        paragraphs: List,
        original: str,
        replacement: str,
    ) -> bool:
        """Ищет original в параграфах и заменяет на replacement."""
        original_clean = self._normalize_text(original)

        for para in paragraphs:
            para_text = self._normalize_text(para.text)
            if not para_text:
                continue

            # Точное совпадение или подстрока
            if original_clean in para_text or self._fuzzy_match(original_clean, para_text):
                # Заменяем текст параграфа, сохраняя форматирование первого run
                self._replace_paragraph_text(para, replacement)
                return True

        return False

    def _find_and_strikethrough(self, paragraphs: List, text_to_remove: str) -> bool:
        """Помечает текст зачёркиванием."""
        text_clean = self._normalize_text(text_to_remove)

        for para in paragraphs:
            para_text = self._normalize_text(para.text)
            if text_clean in para_text or self._fuzzy_match(text_clean, para_text):
                for run in para.runs:
                    run.font.strike = True
                # Добавляем пометку
                strike_run = para.add_run(" [УДАЛЕНО]")
                strike_run.font.color.rgb = RGBColor(255, 0, 0)
                strike_run.bold = True
                return True

        return False

    def _insert_after_section(
        self,
        document: Document,
        paragraphs: List,
        section_title: str,
        text_to_insert: str,
    ) -> bool:
        """Вставляет текст после найденного раздела."""
        if not section_title:
            return False

        section_clean = self._normalize_text(section_title)

        # Ищем заголовок раздела
        target_idx = None
        for i, para in enumerate(paragraphs):
            para_text = self._normalize_text(para.text)
            if section_clean in para_text:
                target_idx = i
                break

        if target_idx is None:
            return False

        # Находим конец раздела (следующий заголовок или конец документа)
        insert_after_idx = target_idx
        for i in range(target_idx + 1, len(paragraphs)):
            para = paragraphs[i]
            # Если это следующий раздел — вставляем перед ним
            if self._is_section_header(para.text):
                break
            insert_after_idx = i

        # Вставляем новый параграф после последнего параграфа раздела
        target_para = paragraphs[insert_after_idx]
        new_para = self._insert_paragraph_after(target_para, text_to_insert)
        if new_para:
            # Помечаем как новый текст
            for run in new_para.runs:
                run.font.color.rgb = RGBColor(0, 100, 0)  # Тёмно-зелёный
            return True

        return False

    def _insert_paragraph_after(self, paragraph, text: str):
        """Вставляет новый параграф после указанного."""
        from docx.oxml.ns import qn
        new_p = paragraph._element.makeelement(qn('w:p'), {})
        paragraph._element.addnext(new_p)
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_p, paragraph._element.getparent())
        run = new_para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return new_para

    def _replace_paragraph_text(self, paragraph, new_text: str):
        """Заменяет текст параграфа, сохраняя базовое форматирование."""
        # Сохраняем форматирование первого run
        first_run_font = None
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_font = {
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'size': first_run.font.size,
                'name': first_run.font.name,
            }

        # Очищаем все runs
        for run in paragraph.runs:
            run.text = ""

        # Вставляем новый текст
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            # Помечаем изменённый текст цветом
            paragraph.runs[0].font.color.rgb = RGBColor(0, 0, 180)  # Синий
        else:
            run = paragraph.add_run(new_text)
            run.font.color.rgb = RGBColor(0, 0, 180)
            if first_run_font:
                run.font.bold = first_run_font.get('bold')
                run.font.size = first_run_font.get('size') or Pt(12)
                run.font.name = first_run_font.get('name') or 'Times New Roman'

    def _normalize_text(self, text: str) -> str:
        """Нормализует текст для поиска."""
        if not text:
            return ""
        # Убираем лишние пробелы и приводим к нижнему регистру
        return re.sub(r'\s+', ' ', text.strip().lower())

    def _fuzzy_match(self, needle: str, haystack: str, threshold: float = 0.7) -> bool:
        """Простой fuzzy match — проверяет, содержит ли haystack достаточную часть needle."""
        if not needle or not haystack:
            return False

        # Разбиваем на слова
        needle_words = set(needle.split())
        haystack_words = set(haystack.split())

        if not needle_words:
            return False

        # Считаем совпадение слов
        common = needle_words & haystack_words
        ratio = len(common) / len(needle_words)
        return ratio >= threshold

    def _is_section_header(self, text: str) -> bool:
        """Проверяет, является ли текст заголовком раздела."""
        stripped = text.strip()
        if not stripped:
            return False
        # Паттерны заголовков
        return bool(re.match(
            r'^(\d+\.?\s+[А-ЯЁA-Z]|[Рр]аздел\s+\d|РАЗДЕЛ\s+\d|Статья\s+\d)',
            stripped
        ))

    def _text_to_docx(self, raw_text: str) -> Document:
        """
        Создаёт DOCX из plain text с базовым форматированием.
        Упрощённая версия TextExtractor._text_to_docx.
        """
        doc = Document()

        # Стиль: Times New Roman 12pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.15

        # Поля страницы (ГОСТ)
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)

        lines = raw_text.split('\n')
        for line in lines:
            stripped = line.strip()
            p = doc.add_paragraph()

            if not stripped:
                continue

            # Заголовок (заглавные)
            is_title = stripped.isupper() and 3 < len(stripped) < 120
            is_section_header = bool(re.match(r'^\d+\.?\s+[А-ЯЁ]', stripped)) and len(stripped) < 120

            if is_title:
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(14)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
            elif is_section_header:
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(13)
                p.paragraph_format.space_before = Pt(10)
            else:
                leading = len(line) - len(line.lstrip())
                run = p.add_run(stripped)
                if leading >= 4:
                    p.paragraph_format.first_line_indent = Cm(1.25)

        return doc

    def _add_recommendation_entry(
        self,
        document: Document,
        idx: int,
        rec: Dict[str, Any],
        applied: bool = True,
    ):
        """Добавляет запись о правке в приложение."""
        section = self._build_section_label(rec)
        action_type = rec.get("action_type", "modify")
        priority = rec.get("priority", "optional")
        reason = rec.get("reason", "")
        original_text = rec.get("original_text", "")
        proposed_text = rec.get("proposed_text", "")

        status = "✅ Применена" if applied else "⚠️ Требует ручного внесения"

        document.add_heading(f"{idx}. {section}", level=2)
        document.add_paragraph(
            f"Статус: {status} | Тип: {self._action_label(action_type)} | "
            f"Приоритет: {priority}"
        )

        if reason:
            document.add_paragraph(f"Обоснование: {reason}")

        if original_text:
            p_old = document.add_paragraph()
            p_old.add_run("Было: ").bold = True
            p_old.add_run(original_text)

        if proposed_text:
            p_new = document.add_paragraph()
            p_new.add_run("Стало: ").bold = True
            p_new.add_run(proposed_text)

    def generate_disagreement_protocol_docx(
        self,
        accepted_recommendations: List[Dict[str, Any]],
        source_file_name: str = "contract",
    ) -> bytes:
        """
        Variant B:
        Generate disagreement protocol DOCX (before -> after table).
        """
        document = Document()
        document.add_heading("Протокол разногласий", level=1)
        document.add_paragraph(f"Документ: {source_file_name}")
        document.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

        if not accepted_recommendations:
            document.add_paragraph("Согласованные разногласия отсутствуют.")
            return self._to_bytes(document)

        table = document.add_table(rows=1, cols=6)
        header_cells = table.rows[0].cells
        header_cells[0].text = "№"
        header_cells[1].text = "Раздел"
        header_cells[2].text = "Тип"
        header_cells[3].text = "Было"
        header_cells[4].text = "Стало"
        header_cells[5].text = "Обоснование"

        for idx, rec in enumerate(accepted_recommendations, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = self._build_section_label(rec)
            row[2].text = self._action_label(rec.get("action_type", "modify"))
            row[3].text = str(rec.get("original_text", ""))
            row[4].text = str(rec.get("proposed_text", ""))
            row[5].text = str(rec.get("reason", ""))

        return self._to_bytes(document)

    def generate_disagreement_protocol_json(
        self,
        accepted_recommendations: List[Dict[str, Any]],
    ) -> str:
        """JSON export for protocol."""
        payload: List[Dict[str, Any]] = []
        for idx, rec in enumerate(accepted_recommendations, 1):
            payload.append(
                {
                    "index": idx,
                    "section": self._build_section_label(rec),
                    "action_type": rec.get("action_type", "modify"),
                    "priority": rec.get("priority", "optional"),
                    "before": rec.get("original_text", ""),
                    "after": rec.get("proposed_text", ""),
                    "reason": rec.get("reason", ""),
                    "source": rec.get("source", "section_analysis"),
                }
            )

        return json.dumps(payload, ensure_ascii=False, indent=2)

    def _to_bytes(self, document: Document) -> bytes:
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _build_section_label(self, rec: Dict[str, Any]) -> str:
        section_number = rec.get("section_number", "")
        section_title = rec.get("section_title", "")
        if section_number and section_title:
            return "{0}. {1}".format(section_number, section_title)
        if section_title:
            return str(section_title)
        if section_number:
            return str(section_number)
        return "Не указан"

    def _action_label(self, action_type: str) -> str:
        mapping = {
            "add": "Добавить",
            "modify": "Изменить",
            "remove": "Удалить",
            "missing": "Добавить отсутствующий пункт",
            "weakened": "Усилить условие",
            "contradicts": "Исправить противоречие",
            "added": "Проверить добавленное условие",
        }
        return mapping.get(str(action_type).lower(), "Изменить")
