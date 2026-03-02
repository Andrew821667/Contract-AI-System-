# -*- coding: utf-8 -*-
"""
Сервис генерации договоров через LLM
Простой LLM-based сервис: тип + параметры → текст договора → DOCX
"""
import os
import json
from typing import Dict, Any, Optional, List
from datetime import datetime
from dataclasses import dataclass, field
from loguru import logger


@dataclass
class ContractParty:
    """Сторона договора."""
    name: str
    inn: str = ""
    ogrn: str = ""
    address: str = ""
    representative: str = ""
    position: str = ""
    basis: str = "Устава"


@dataclass
class ContractParams:
    """Параметры генерируемого договора."""
    contract_type: str  # код из contract_types.py
    party_a: ContractParty = field(default_factory=lambda: ContractParty(name=""))
    party_b: ContractParty = field(default_factory=lambda: ContractParty(name=""))
    subject: str = ""
    amount: str = ""
    currency: str = "рублей"
    duration: str = ""
    start_date: str = ""
    payment_terms: str = ""
    additional_conditions: str = ""
    city: str = "Москва"


@dataclass
class GenerationResult:
    """Результат генерации договора."""
    success: bool
    contract_text: str = ""
    docx_path: str = ""
    error: str = ""
    tokens_used: int = 0
    generation_time: float = 0.0
    contract_type: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)


class ContractGenerationService:
    """
    Генерация текста договора через LLM + форматирование в DOCX.

    Использование:
        service = ContractGenerationService()
        params = ContractParams(
            contract_type="supply",
            party_a=ContractParty(name="ООО Альфа", inn="7701234567"),
            party_b=ContractParty(name="ООО Бета", inn="7709876543"),
            subject="товары бытовой химии",
            amount="1 500 000",
            duration="12 месяцев",
        )
        result = service.generate(params)
    """

    SYSTEM_PROMPT = """Ты — опытный юрист-составитель договоров по праву Российской Федерации.

ЗАДАЧА: Составить полный текст договора на основе предоставленных параметров.

ТРЕБОВАНИЯ:
1. Договор должен соответствовать нормам ГК РФ для данного типа договора
2. Включи все существенные условия, необходимые по закону
3. Используй профессиональный юридический язык
4. Структура: преамбула, предмет, права и обязанности сторон, цена и порядок расчётов,
   ответственность, форс-мажор, конфиденциальность, срок действия, разрешение споров,
   заключительные положения, реквизиты сторон
5. Нумерация пунктов: 1., 1.1., 1.1.1. и т.д.
6. Пропиши конкретные суммы, сроки и условия из параметров
7. Если каких-то данных нет — укажи [ЗАПОЛНИТЬ] в соответствующих местах

ФОРМАТ ОТВЕТА: Только текст договора, без комментариев и пояснений.
Начни с заголовка и номера договора."""

    def __init__(self, provider: Optional[str] = None, exports_dir: str = "data/exports/contracts"):
        self.provider = provider
        self.exports_dir = exports_dir
        os.makedirs(self.exports_dir, exist_ok=True)

    def generate(self, params: ContractParams) -> GenerationResult:
        """Генерирует текст договора и сохраняет в DOCX."""
        import time
        start_time = time.time()

        try:
            # Получить русское название типа
            from src.utils.contract_types import get_contract_type_name
            type_name = get_contract_type_name(params.contract_type)

            # Сформировать промпт
            prompt = self._build_prompt(params, type_name)
            logger.info(f"Генерация договора: {type_name} | {params.party_a.name} ↔ {params.party_b.name}")

            # Вызвать LLM
            from src.services.llm_gateway import LLMGateway
            gateway = LLMGateway(provider=self.provider)

            contract_text = gateway.call(
                prompt=prompt,
                system_prompt=self.SYSTEM_PROMPT,
                response_format="text",
                temperature=0.3,
                max_tokens=8000,
            )

            if not contract_text or len(contract_text) < 200:
                return GenerationResult(
                    success=False,
                    error="LLM вернула слишком короткий ответ"
                )

            # Сформировать DOCX
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{params.contract_type}_{timestamp}.docx"
            docx_path = os.path.join(self.exports_dir, filename)
            self._save_to_docx(contract_text, type_name, docx_path)

            elapsed = time.time() - start_time
            logger.info(f"Договор сгенерирован за {elapsed:.1f}с: {docx_path}")

            return GenerationResult(
                success=True,
                contract_text=contract_text,
                docx_path=docx_path,
                tokens_used=gateway.total_input_tokens + gateway.total_output_tokens,
                generation_time=elapsed,
                contract_type=params.contract_type,
                metadata={
                    "type_name": type_name,
                    "party_a": params.party_a.name,
                    "party_b": params.party_b.name,
                    "amount": params.amount,
                    "generated_at": datetime.now().isoformat(),
                },
            )

        except Exception as e:
            logger.error(f"Ошибка генерации договора: {e}")
            return GenerationResult(success=False, error=str(e))

    def _build_prompt(self, params: ContractParams, type_name: str) -> str:
        """Формирует промпт для LLM из параметров."""
        parts = [f"Составь {type_name.upper()}\n"]

        parts.append("ДАННЫЕ СТОРОН:")
        parts.append(self._format_party(params.party_a, "Сторона 1"))
        parts.append(self._format_party(params.party_b, "Сторона 2"))

        parts.append("\nПАРАМЕТРЫ ДОГОВОРА:")
        if params.subject:
            parts.append(f"- Предмет: {params.subject}")
        if params.amount:
            parts.append(f"- Сумма: {params.amount} {params.currency}")
        if params.duration:
            parts.append(f"- Срок: {params.duration}")
        if params.start_date:
            parts.append(f"- Дата начала: {params.start_date}")
        if params.payment_terms:
            parts.append(f"- Условия оплаты: {params.payment_terms}")
        if params.city:
            parts.append(f"- Город заключения: {params.city}")
        if params.additional_conditions:
            parts.append(f"\nДОПОЛНИТЕЛЬНЫЕ УСЛОВИЯ:\n{params.additional_conditions}")

        return "\n".join(parts)

    def _format_party(self, party: ContractParty, label: str) -> str:
        """Форматирует данные стороны для промпта."""
        lines = [f"\n{label}: {party.name}"]
        if party.inn:
            lines.append(f"  ИНН: {party.inn}")
        if party.ogrn:
            lines.append(f"  ОГРН: {party.ogrn}")
        if party.address:
            lines.append(f"  Адрес: {party.address}")
        if party.representative:
            lines.append(f"  Представитель: {party.representative}")
        if party.position:
            lines.append(f"  Должность: {party.position}")
        if party.basis:
            lines.append(f"  Действует на основании: {party.basis}")
        return "\n".join(lines)

    def _save_to_docx(self, text: str, title: str, file_path: str):
        """Сохраняет текст договора в DOCX с форматированием."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Шрифт по умолчанию
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.15

        # Поля страницы
        for section in doc.sections:
            section.left_margin = Inches(1.18)
            section.right_margin = Inches(0.59)
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.79)

        # Разбиваем текст на строки и форматируем
        lines = text.strip().split("\n")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            # Заголовок договора (первые строки, заглавные или содержат "ДОГОВОР")
            if stripped.isupper() and len(stripped) < 100:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(14)
                continue

            # Заголовки разделов (начинаются с цифры и точки, полностью заглавные)
            if stripped and stripped[0].isdigit() and "." in stripped[:5]:
                # Проверяем: это заголовок раздела или обычный пункт?
                after_number = stripped.split(".", 1)[-1].strip() if "." in stripped else ""
                if after_number.isupper() or (after_number and not any(c.islower() for c in after_number[:20])):
                    p = doc.add_paragraph()
                    run = p.add_run(stripped)
                    run.bold = True
                    run.font.size = Pt(12)
                    continue

            # Обычный текст
            p = doc.add_paragraph(stripped)
            p.paragraph_format.first_line_indent = Inches(0.49)

        doc.save(file_path)
        logger.info(f"DOCX сохранён: {file_path}")


__all__ = ["ContractGenerationService", "ContractParams", "ContractParty", "GenerationResult"]
