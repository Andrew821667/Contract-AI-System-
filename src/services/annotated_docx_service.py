# -*- coding: utf-8 -*-
"""
Annotated DOCX Export Service
Creates DOCX with highlighted risks and comments
"""
import io
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger


class AnnotatedDocxService:
    """Generates annotated DOCX documents with risk highlighting"""

    # Risk severity → background color
    SEVERITY_COLORS = {
        "critical": "FF9999",   # Red
        "significant": "FF9999",  # Red
        "high": "FF9999",      # Red
        "medium": "FFFF99",    # Yellow
        "minor": "CCFFCC",     # Light green
        "low": "CCFFCC",       # Light green
    }

    def create_annotated_docx(self, contract, analysis, db) -> bytes:
        """
        Create an annotated DOCX with risk highlights and comments.

        Args:
            contract: Contract model instance
            analysis: AnalysisResult model instance
            db: Database session

        Returns:
            DOCX file bytes
        """
        doc = Document()

        # Title
        title = doc.add_heading(f"Аннотированный анализ: {contract.file_name or 'Договор'}", level=0)

        # Summary section
        doc.add_heading("Сводка анализа", level=1)

        risks = []
        recommendations = []

        # Extract risks from analysis
        if hasattr(analysis, 'risks_by_category') and analysis.risks_by_category:
            if isinstance(analysis.risks_by_category, list):
                risks = analysis.risks_by_category
            elif isinstance(analysis.risks_by_category, dict):
                for category, category_risks in analysis.risks_by_category.items():
                    if isinstance(category_risks, list):
                        risks.extend(category_risks)

        if hasattr(analysis, 'recommendations') and analysis.recommendations:
            if isinstance(analysis.recommendations, list):
                recommendations = analysis.recommendations

        # Stats paragraph
        stats = doc.add_paragraph()
        stats.add_run(f"Всего рисков: {len(risks)}").bold = True
        stats.add_run(f" | Рекомендаций: {len(recommendations)}")

        # Risk table
        if risks:
            doc.add_heading("Выявленные риски", level=1)

            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Header
            headers = ['Риск', 'Серьёзность', 'Вероятность', 'Описание']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Rows
            for risk in risks:
                if isinstance(risk, dict):
                    row = table.add_row()
                    title_text = risk.get('title', risk.get('risk_type', 'N/A'))
                    severity = risk.get('severity', 'medium')
                    probability = risk.get('probability', 'N/A')
                    description = risk.get('description', '')

                    row.cells[0].text = str(title_text)
                    row.cells[1].text = str(severity)
                    row.cells[2].text = str(probability)
                    row.cells[3].text = str(description)[:200]

                    # Highlight severity cell
                    color = self.SEVERITY_COLORS.get(str(severity).lower(), "FFFFFF")
                    self._set_cell_bg(row.cells[1], color)

        # Recommendations
        if recommendations:
            doc.add_heading("Рекомендации", level=1)
            for idx, rec in enumerate(recommendations, 1):
                if isinstance(rec, dict):
                    title = rec.get('title', f'Рекомендация {idx}')
                    priority = rec.get('priority', 'medium')
                    description = rec.get('description', '')

                    p = doc.add_paragraph()
                    run = p.add_run(f"{idx}. [{priority.upper()}] {title}")
                    run.bold = True

                    if description:
                        doc.add_paragraph(description, style='List Bullet')
                elif isinstance(rec, str):
                    doc.add_paragraph(f"{idx}. {rec}", style='List Bullet')

        # Contract text with annotations (if original file is DOCX)
        if contract.file_path and contract.file_path.endswith('.docx'):
            try:
                doc.add_page_break()
                doc.add_heading("Текст договора с аннотациями", level=1)

                original = Document(contract.file_path)
                for para in original.paragraphs:
                    if para.text.strip():
                        new_para = doc.add_paragraph()

                        # Check if any risk references this text
                        matched_risk = self._find_matching_risk(para.text, risks)

                        if matched_risk:
                            severity = matched_risk.get('severity', 'medium')
                            color = self.SEVERITY_COLORS.get(str(severity).lower(), "FFFFFF")

                            run = new_para.add_run(para.text)
                            self._set_run_highlight(run, color)

                            # Add comment-like annotation
                            comment = new_para.add_run(
                                f"  [!] {matched_risk.get('title', 'Риск')}: {matched_risk.get('description', '')[:100]}"
                            )
                            comment.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                            comment.font.size = Pt(8)
                            comment.italic = True
                        else:
                            new_para.add_run(para.text)
            except Exception as e:
                logger.warning(f"Could not annotate original DOCX: {e}")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _find_matching_risk(self, text: str, risks: list) -> Optional[dict]:
        """Find if any risk matches the given text"""
        text_lower = text.lower()
        for risk in risks:
            if not isinstance(risk, dict):
                continue
            # Match by keywords from risk title or description
            title = str(risk.get('title', '')).lower()
            if title and len(title) > 5:
                # Check if key words overlap
                title_words = set(title.split())
                text_words = set(text_lower.split())
                overlap = title_words & text_words
                if len(overlap) >= 2:
                    return risk
        return None

    @staticmethod
    def _set_cell_bg(cell, color_hex: str):
        """Set cell background color"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _set_run_highlight(run, color_hex: str):
        """Set run background highlight"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        shading.set(qn('w:val'), 'clear')
        run._r.get_or_add_rPr().append(shading)
