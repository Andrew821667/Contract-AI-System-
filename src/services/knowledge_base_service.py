# -*- coding: utf-8 -*-
"""
Knowledge Base Service — управление RAG базой знаний

Предоставляет CRUD-операции для документов в PostgreSQL и ChromaDB,
загрузку файлов с дедупликацией по SHA256, чанкование и индексацию.
"""
import hashlib
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Dict, Any, Optional, Callable

from loguru import logger
from sqlalchemy import func, desc
from sqlalchemy.orm import Session

from ..models.database import LegalDocument


class KnowledgeBaseService:
    """Сервис управления базой знаний RAG"""

    SECTIONS = {
        'law': {'label': 'Нормативная база РФ', 'collection': 'laws', 'icon': '📜'},
        'case_law': {'label': 'Судебная практика РФ', 'collection': 'case_law', 'icon': '⚖️'},
        'company_kb': {'label': 'БЗ компании', 'collection': 'company_kb', 'icon': '🏢'},
        'template': {'label': 'Образцы документов', 'collection': 'templates', 'icon': '📄'},
    }

    def __init__(self, db: Session, rag_system=None):
        """
        Args:
            db: SQLAlchemy session
            rag_system: RAGSystem instance (optional, for vectorization)
        """
        self.db = db
        self.rag_system = rag_system

    # ─── Просмотр ───────────────────────────────────────────────

    def get_documents(
        self,
        doc_type: str,
        search_query: Optional[str] = None,
        limit: int = 100,
        offset: int = 0
    ) -> List[LegalDocument]:
        """Получить документы по типу с опциональным поиском по названию"""
        query = self.db.query(LegalDocument).filter(
            LegalDocument.doc_type == doc_type,
            LegalDocument.status.in_(['active', 'pending', 'processing', 'error'])
        )

        if search_query:
            safe_q = search_query.replace('%', r'\%').replace('_', r'\_')
            query = query.filter(
                LegalDocument.title.ilike(f'%{safe_q}%', escape='\\')
            )

        return query.order_by(desc(LegalDocument.created_at)).offset(offset).limit(limit).all()

    def get_document_by_id(self, doc_id: str) -> Optional[LegalDocument]:
        """Получить документ по ID (primary key)"""
        return self.db.query(LegalDocument).filter(LegalDocument.id == doc_id).first()

    def get_document_by_doc_id(self, doc_id: str) -> Optional[LegalDocument]:
        """Получить документ по doc_id"""
        return self.db.query(LegalDocument).filter(LegalDocument.doc_id == doc_id).first()

    def get_section_stats(self, doc_type: str) -> Dict[str, Any]:
        """Статистика раздела: кол-во документов, чанков, последнее обновление"""
        docs = self.db.query(LegalDocument).filter(
            LegalDocument.doc_type == doc_type,
            LegalDocument.status.in_(['active', 'pending', 'processing', 'error'])
        )

        doc_count = docs.count()

        chunks_total = self.db.query(func.sum(LegalDocument.chunks_count)).filter(
            LegalDocument.doc_type == doc_type,
            LegalDocument.status == 'active'
        ).scalar() or 0

        last_updated = self.db.query(func.max(LegalDocument.updated_at)).filter(
            LegalDocument.doc_type == doc_type
        ).scalar()

        vectorized_count = self.db.query(LegalDocument).filter(
            LegalDocument.doc_type == doc_type,
            LegalDocument.is_vectorized == True,
            LegalDocument.status == 'active'
        ).count()

        return {
            'count': doc_count,
            'chunks': chunks_total,
            'last_updated': last_updated,
            'vectorized': vectorized_count,
        }

    # ─── Редактирование ─────────────────────────────────────────

    def update_document(self, doc_id: str, new_content: str, new_title: Optional[str] = None) -> LegalDocument:
        """Обновить содержимое документа (без переиндексации)"""
        doc = self.db.query(LegalDocument).filter(LegalDocument.id == doc_id).first()
        if not doc:
            raise ValueError(f"Документ {doc_id} не найден")

        doc.content = new_content
        if new_title is not None:
            doc.title = new_title
        doc.is_vectorized = False
        doc.updated_at = datetime.now(timezone.utc)

        self.db.commit()
        self.db.refresh(doc)
        logger.info(f"Документ {doc.doc_id} обновлён")
        return doc

    def delete_document(self, doc_id: str) -> bool:
        """Удалить документ из PostgreSQL и ChromaDB"""
        doc = self.db.query(LegalDocument).filter(LegalDocument.id == doc_id).first()
        if not doc:
            return False

        # Удаление из ChromaDB
        if self.rag_system:
            section = self.SECTIONS.get(doc.doc_type)
            collection_name = section['collection'] if section else 'knowledge'
            try:
                self.rag_system.delete_document(doc.doc_id, collection_name)
            except Exception as e:
                logger.warning(f"Ошибка удаления из ChromaDB: {e}")

        # Удаление файла
        if doc.file_path and os.path.exists(doc.file_path):
            try:
                os.remove(doc.file_path)
            except OSError as e:
                logger.warning(f"Не удалось удалить файл {doc.file_path}: {e}")

        self.db.delete(doc)
        self.db.commit()
        logger.info(f"Документ {doc.doc_id} удалён")
        return True

    def reindex_document(self, doc_id: str) -> int:
        """Переиндексировать документ: перечанковать + перевекторизовать"""
        doc = self.db.query(LegalDocument).filter(LegalDocument.id == doc_id).first()
        if not doc:
            raise ValueError(f"Документ {doc_id} не найден")

        if not self.rag_system:
            raise RuntimeError("RAG система не инициализирована")

        section = self.SECTIONS.get(doc.doc_type)
        collection_name = section['collection'] if section else 'knowledge'

        # Удалить старые чанки
        try:
            self.rag_system.delete_document(doc.doc_id, collection_name)
        except Exception as e:
            logger.warning(f"Ошибка удаления старых чанков: {e}")

        # Индексировать заново
        metadata = {
            'title': doc.title,
            'doc_type': doc.doc_type,
            'source': doc.source or 'manual',
        }

        self.rag_system.index_document(
            doc_id=doc.doc_id,
            content=doc.content,
            metadata=metadata,
            collection=collection_name
        )

        # Подсчитать чанки
        try:
            coll = self.rag_system.collections.get(collection_name)
            if coll:
                results = coll.get(where={"doc_id": doc.doc_id})
                chunks_count = len(results['ids']) if results['ids'] else 0
            else:
                chunks_count = 0
        except Exception:
            chunks_count = 0

        doc.is_vectorized = True
        doc.chunks_count = chunks_count
        doc.updated_at = datetime.now(timezone.utc)
        self.db.commit()

        logger.info(f"Документ {doc.doc_id} переиндексирован: {chunks_count} чанков")
        return chunks_count

    # ─── Загрузка ────────────────────────────────────────────────

    def compute_file_hash(self, file_data: bytes) -> str:
        """SHA256 хеш файла"""
        return hashlib.sha256(file_data).hexdigest()

    def find_duplicate(self, file_hash: str, doc_type: str) -> Optional[LegalDocument]:
        """Найти дубликат по хешу файла в рамках типа"""
        return self.db.query(LegalDocument).filter(
            LegalDocument.file_hash == file_hash,
            LegalDocument.doc_type == doc_type,
            LegalDocument.status.in_(['active', 'pending', 'processing'])
        ).first()

    def convert_to_markdown(self, file_data: bytes, filename: str) -> str:
        """Конвертировать файл в текст/markdown"""
        ext = Path(filename).suffix.lower()

        if ext in ('.txt', '.md'):
            # Попытка декодировать как UTF-8, иначе cp1251
            try:
                return file_data.decode('utf-8')
            except UnicodeDecodeError:
                return file_data.decode('cp1251', errors='replace')

        elif ext == '.docx':
            return self._convert_docx(file_data)

        elif ext == '.pdf':
            return self._convert_pdf(file_data)

        else:
            raise ValueError(f"Неподдерживаемый формат: {ext}")

    def _convert_docx(self, file_data: bytes) -> str:
        """Конвертация DOCX → текст"""
        import io
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_data))
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Заголовки → markdown headers
                    if para.style and para.style.name.startswith('Heading'):
                        level = 1
                        try:
                            level = int(para.style.name.split()[-1])
                        except (ValueError, IndexError):
                            level = 1
                        paragraphs.append(f"{'#' * level} {text}")
                    else:
                        paragraphs.append(text)
            return '\n\n'.join(paragraphs)
        except Exception as e:
            logger.error(f"Ошибка конвертации DOCX: {e}")
            raise ValueError(f"Не удалось прочитать DOCX: {e}")

    def _convert_pdf(self, file_data: bytes) -> str:
        """Конвертация PDF → текст"""
        import io
        try:
            import pdfplumber
            pages_text = []
            with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
            if pages_text:
                return '\n\n'.join(pages_text)
        except Exception as e:
            logger.warning(f"pdfplumber не сработал: {e}")

        # Fallback: pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_data))
            pages_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            return '\n\n'.join(pages_text)
        except Exception as e:
            logger.error(f"Ошибка конвертации PDF: {e}")
            raise ValueError(f"Не удалось прочитать PDF: {e}")

    def process_upload(
        self,
        file_data: bytes,
        filename: str,
        doc_type: str,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Обработка одного загруженного файла.

        Returns:
            {status: 'created'|'updated'|'duplicate', doc_id: str, chunks_count: int}
        """
        file_hash = self.compute_file_hash(file_data)

        # Проверка дубликата
        existing = self.find_duplicate(file_hash, doc_type)
        if existing:
            return {
                'status': 'duplicate',
                'doc_id': existing.doc_id,
                'chunks_count': existing.chunks_count,
                'title': existing.title,
            }

        # Конвертация
        content = self.convert_to_markdown(file_data, filename)
        if not content.strip():
            raise ValueError(f"Файл {filename} пуст или не удалось извлечь текст")

        doc_title = title or Path(filename).stem.replace('_', ' ').replace('-', ' ')
        doc_id = f"{doc_type}_{uuid.uuid4().hex[:12]}"

        # Проверка: тот же doc_id существует (обновление)
        # Так как doc_id генерируется, коллизии маловероятны

        # Сохранение файла
        upload_dir = Path("data/knowledge_base") / doc_type
        upload_dir.mkdir(parents=True, exist_ok=True)
        safe_name = f"{doc_id}_{Path(filename).name}"
        file_path = str(upload_dir / safe_name)

        with open(file_path, 'wb') as f:
            f.write(file_data)

        # Создание записи в БД
        doc = LegalDocument(
            doc_id=doc_id,
            title=doc_title,
            doc_type=doc_type,
            content=content,
            status='active',
            is_vectorized=False,
            file_hash=file_hash,
            file_name=filename,
            file_path=file_path,
            source='manual',
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        # Индексация в ChromaDB
        chunks_count = 0
        if self.rag_system:
            try:
                chunks_count = self.reindex_document(doc.id)
            except Exception as e:
                logger.error(f"Ошибка индексации {doc_id}: {e}")
                doc.status = 'error'
                self.db.commit()
                return {
                    'status': 'error',
                    'doc_id': doc_id,
                    'chunks_count': 0,
                    'title': doc_title,
                    'error': str(e),
                }

        return {
            'status': 'created',
            'doc_id': doc_id,
            'chunks_count': chunks_count,
            'title': doc_title,
        }

    def process_batch(
        self,
        files: List[Dict[str, Any]],
        doc_type: str,
        on_progress: Optional[Callable[[int, int, str], None]] = None
    ) -> List[Dict[str, Any]]:
        """
        Обработка пакета файлов.

        Args:
            files: List of {data: bytes, name: str}
            doc_type: Тип документа
            on_progress: callback(current, total, message)

        Returns:
            List of results per file
        """
        results = []
        total = len(files)

        for i, file_info in enumerate(files):
            filename = file_info['name']
            file_data = file_info['data']

            if on_progress:
                on_progress(i, total, f"Обработка: {filename}")

            try:
                result = self.process_upload(file_data, filename, doc_type)
                result['filename'] = filename
                results.append(result)
            except Exception as e:
                logger.error(f"Ошибка обработки {filename}: {e}")
                results.append({
                    'status': 'error',
                    'filename': filename,
                    'doc_id': None,
                    'chunks_count': 0,
                    'error': str(e),
                })

        if on_progress:
            on_progress(total, total, "Готово")

        return results

    # ─── Массовые операции ───────────────────────────────────────

    def reindex_all(self, doc_type: Optional[str] = None) -> int:
        """Переиндексировать все документы (или по типу)"""
        query = self.db.query(LegalDocument).filter(
            LegalDocument.status == 'active'
        )
        if doc_type:
            query = query.filter(LegalDocument.doc_type == doc_type)

        docs = query.all()
        count = 0
        for doc in docs:
            try:
                self.reindex_document(doc.id)
                count += 1
            except Exception as e:
                logger.error(f"Ошибка переиндексации {doc.doc_id}: {e}")

        logger.info(f"Переиндексировано {count}/{len(docs)} документов")
        return count

    def reindex_pending(self) -> int:
        """Переиндексировать только документы с is_vectorized=False"""
        docs = self.db.query(LegalDocument).filter(
            LegalDocument.is_vectorized == False,
            LegalDocument.status == 'active'
        ).all()

        count = 0
        for doc in docs:
            try:
                self.reindex_document(doc.id)
                count += 1
            except Exception as e:
                logger.error(f"Ошибка переиндексации {doc.doc_id}: {e}")

        logger.info(f"Переиндексировано pending: {count}/{len(docs)}")
        return count
