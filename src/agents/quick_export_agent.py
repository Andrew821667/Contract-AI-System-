# -*- coding: utf-8 -*-
"""
Quick Export Agent - Fast export of contracts to various formats
"""
import os
import json
import re
import shutil
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime, timezone
from loguru import logger

from .base_agent import BaseAgent, AgentResult
from ..services.llm_gateway import LLMGateway
from ..services.document_parser_extended import ExtendedDocumentParser
from ..models.database import Contract, ExportLog
from ..utils.xml_security import parse_xml_safely


class QuickExportAgent(BaseAgent):
    """
    Agent for quick export of contracts

    Capabilities:
    - Export original DOCX/PDF без потери форматирования
    - Export to plain text
    - Export canonical XML
    - Export metadata (JSON)
    - Batch export
    - Export logging
    """

    def __init__(
        self,
        llm_gateway: LLMGateway,
        db_session,
        config: Optional[Dict[str, Any]] = None
    ):
        super().__init__(llm_gateway, db_session, config)
        self.export_dir = self.config.get('export_dir', 'data/exports')
        os.makedirs(self.export_dir, exist_ok=True)

    def get_name(self) -> str:
        return "QuickExportAgent"

    def get_system_prompt(self) -> str:
        return "You are a document export assistant."

    def execute(self, state: Dict[str, Any]) -> AgentResult:
        """
        Execute quick export

        Expected state:
        - contract_id: ID of contract to export
        - export_format: 'docx', 'pdf', 'txt', 'json', 'xml', 'all'
        - include_analysis: Include analysis results (default: False)
        - allow_lossy_conversion: Allow best-effort cross-format conversion
        - user_id: ID of user requesting export

        Returns:
        - file_paths: Dict of exported file paths
        - export_log_id: ID of export log record
        """
        try:
            contract_id = state.get('contract_id')
            export_format = state.get('export_format', 'docx')
            include_analysis = state.get('include_analysis', False)
            allow_lossy_conversion = state.get('allow_lossy_conversion', False)
            user_id = state.get('user_id')

            if not contract_id:
                return AgentResult(
                    success=False,
                    data={},
                    error="Missing contract_id"
                )

            logger.info(f"Quick export: contract {contract_id}, format {export_format}")

            # Get contract
            contract = self.db.query(Contract).filter(
                Contract.id == contract_id
            ).first()

            if not contract:
                return AgentResult(
                    success=False,
                    data={},
                    error=f"Contract {contract_id} not found"
                )

            # Export based on format
            file_paths = {}

            if export_format == 'all':
                formats = ['docx', 'pdf', 'txt', 'json', 'xml']
            else:
                formats = [export_format]

            for fmt in formats:
                try:
                    file_path = self._export_format(contract, fmt, include_analysis, allow_lossy_conversion)
                    file_paths[fmt] = file_path
                except Exception as e:
                    logger.error(f"Failed to export {fmt}: {e}")
                    file_paths[fmt] = None

            # Log export
            export_log = self._log_export(contract_id, export_format, file_paths, user_id)

            logger.info(f"Export complete: {len(file_paths)} files")

            return AgentResult(
                success=True,
                data={
                    'contract_id': contract_id,
                    'file_paths': file_paths,
                    'export_log_id': export_log.id if export_log else None
                },
                metadata={'message': f"Exported to {len(file_paths)} format(s)"}
            )

        except Exception as e:
            logger.error(f"Quick export failed: {e}")
            return AgentResult(
                success=False,
                data={},
                error=str(e)
            )

    def _export_format(
        self,
        contract: Contract,
        format: str,
        include_analysis: bool,
        allow_lossy_conversion: bool,
    ) -> str:
        """Export contract to specific format"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_stem = Path(contract.file_name or f"contract_{contract.id}").stem
        safe_stem = safe_stem.replace('/', '_').replace('\\', '_')
        base_name = f"{safe_stem}_{timestamp}"

        if format == 'docx':
            return self._export_docx(contract, base_name, include_analysis, allow_lossy_conversion)
        elif format == 'pdf':
            return self._export_pdf(contract, base_name, include_analysis, allow_lossy_conversion)
        elif format == 'txt':
            return self._export_txt(contract, base_name, include_analysis)
        elif format == 'json':
            return self._export_json(contract, base_name, include_analysis)
        elif format == 'xml':
            return self._export_xml(contract, base_name, include_analysis)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _export_docx(
        self,
        contract: Contract,
        base_name: str,
        include_analysis: bool,
        allow_lossy_conversion: bool = False,
    ) -> str:
        """Export DOCX preserving original formatting whenever possible."""
        output_path = os.path.join(self.export_dir, f"{base_name}.docx")
        original_ext = self._get_original_extension(contract)

        if original_ext == '.docx' and not include_analysis:
            shutil.copy2(contract.file_path, output_path)
            logger.info(f"DOCX export completed by copying original file: {output_path}")
            return output_path

        if include_analysis and original_ext == '.docx':
            from ..models.database import AnalysisResult
            from ..services.annotated_docx_service import AnnotatedDocxService

            analysis = (
                self.db.query(AnalysisResult)
                .filter(AnalysisResult.contract_id == contract.id)
                .order_by(AnalysisResult.created_at.desc(), AnalysisResult.version.desc())
                .first()
            )
            if not analysis:
                raise ValueError("Не найден результат анализа для аннотированного DOCX")

            service = AnnotatedDocxService()
            docx_bytes = service.create_annotated_docx(contract, analysis, self.db)
            with open(output_path, 'wb') as f:
                f.write(docx_bytes)
            logger.info(f"Annotated DOCX export completed: {output_path}")
            return output_path

        if allow_lossy_conversion:
            xml_content = self._get_canonical_xml(contract)
            self._write_reconstructed_docx(contract, output_path, xml_content)
            logger.info(f"Best-effort DOCX conversion completed: {output_path}")
            return output_path

        raise ValueError(
            "Экспорт в DOCX без потери форматирования доступен только для исходных DOCX-файлов. "
            "Для кросс-конверсии подтвердите предупреждение о возможной потере форматирования."
        )

    def _export_pdf(
        self,
        contract: Contract,
        base_name: str,
        include_analysis: bool,
        allow_lossy_conversion: bool = False,
    ) -> str:
        """Export PDF preserving original formatting whenever possible."""
        output_path = os.path.join(self.export_dir, f"{base_name}.pdf")
        original_ext = self._get_original_extension(contract)
        if original_ext == '.pdf' and not include_analysis:
            shutil.copy2(contract.file_path, output_path)
            logger.info(f"PDF export completed by copying original file: {output_path}")
            return output_path

        if allow_lossy_conversion:
            xml_content = self._get_canonical_xml(contract)
            self._write_reconstructed_pdf(contract, output_path, xml_content)
            logger.info(f"Best-effort PDF conversion completed: {output_path}")
            return output_path

        raise ValueError(
            "Экспорт в PDF без потери форматирования доступен только для исходных PDF-файлов. "
            "Для кросс-конверсии подтвердите предупреждение о возможной потере форматирования."
        )

    def _export_txt(self, contract: Contract, base_name: str, include_analysis: bool) -> str:
        """Export to plain text"""
        output_path = os.path.join(self.export_dir, f"{base_name}.txt")
        original_ext = self._get_original_extension(contract)
        if original_ext == '.txt' and not include_analysis:
            shutil.copy2(contract.file_path, output_path)
            logger.info(f"TXT export completed by copying original file: {output_path}")
            return output_path

        xml_content = self._get_canonical_xml(contract)
        content: List[str] = [self._xml_to_plain_text(xml_content)]

        if include_analysis and contract.analysis_results:
            latest = sorted(
                contract.analysis_results,
                key=lambda item: (item.created_at or datetime(1970, 1, 1, tzinfo=timezone.utc), item.version or 0),
                reverse=True,
            )[0]
            content.extend([
                "",
                "",
                "=" * 80,
                "РЕЗУЛЬТАТЫ АНАЛИЗА",
                "=" * 80,
                json.dumps(
                    {
                        'analysis_id': latest.id,
                        'version': latest.version,
                        'created_at': latest.created_at.isoformat() if latest.created_at else None,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
            ])
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))

        logger.info(f"TXT export: {output_path}")
        return output_path

    def _export_json(self, contract: Contract, base_name: str, include_analysis: bool) -> str:
        """Export metadata to JSON"""
        import json
        output_path = os.path.join(self.export_dir, f"{base_name}.json")

        data = {
            'id': contract.id,
            'file_name': contract.file_name,
            'file_path': contract.file_path,
            'document_type': contract.document_type,
            'contract_type': contract.contract_type,
            'upload_date': contract.upload_date.isoformat() if contract.upload_date else None,
            'status': contract.status,
            'risk_level': contract.risk_level,
            'meta_info': contract.meta_info
        }

        if include_analysis and contract.analysis_results:
            latest = sorted(
                contract.analysis_results,
                key=lambda item: (item.created_at or datetime(1970, 1, 1, tzinfo=timezone.utc), item.version or 0),
                reverse=True,
            )[0]
            data['analysis_results'] = [
                {
                    'id': a.id,
                    'version': a.version,
                    'created_at': a.created_at.isoformat() if a.created_at else None
                }
                for a in contract.analysis_results
            ]
            data['latest_analysis'] = {
                'id': latest.id,
                'version': latest.version,
                'entities': latest.entities,
                'compliance_issues': latest.compliance_issues,
                'legal_issues': latest.legal_issues,
                'risks_by_category': latest.risks_by_category,
                'recommendations': latest.recommendations,
            }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        logger.info(f"JSON export: {output_path}")
        return output_path

    def _export_xml(self, contract: Contract, base_name: str, include_analysis: bool) -> str:
        """Export canonical XML used by the parser/analysis pipeline."""
        output_path = os.path.join(self.export_dir, f"{base_name}.xml")
        original_ext = self._get_original_extension(contract)

        if original_ext == '.xml' and not include_analysis:
            shutil.copy2(contract.file_path, output_path)
            logger.info(f"XML export completed by copying original file: {output_path}")
            return output_path

        xml_content = self._get_canonical_xml(contract)
        parse_xml_safely(xml_content)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(xml_content)

        logger.info(f"XML export completed: {output_path}")
        return output_path

    @staticmethod
    def _get_original_extension(contract: Contract) -> str:
        return Path(contract.file_name or contract.file_path or '').suffix.lower()

    def _get_canonical_xml(self, contract: Contract) -> str:
        meta_info = contract.meta_info if isinstance(contract.meta_info, dict) else {}
        cached_xml = meta_info.get('xml') if isinstance(meta_info, dict) else None
        if isinstance(cached_xml, str) and cached_xml.strip():
            return cached_xml

        parser = ExtendedDocumentParser()
        xml_content = parser.parse(contract.file_path)
        return xml_content if isinstance(xml_content, str) else str(xml_content)

    @staticmethod
    def _looks_like_document_title(text: str) -> bool:
        normalized = " ".join(text.split())
        if not normalized or len(normalized) > 200:
            return False
        return bool(re.match(r'^(договор|соглашение|контракт|акт|протокол)\b', normalized, re.IGNORECASE))

    @classmethod
    def _extract_structured_content(cls, xml_content: str) -> List[Dict[str, Any]]:
        root = parse_xml_safely(xml_content)
        items: List[Dict[str, Any]] = []
        title_added = False

        document_title = (root.findtext('./metadata/title', '') or '').strip()
        if document_title and document_title.lower() not in {'без названия', 'untitled'}:
            items.append({'kind': 'title', 'text': document_title})
            title_added = True

        for clause in root.findall('./clauses/clause'):
            title = (clause.findtext('title', '') or '').strip()
            if title and not (clause.get('type') == 'preamble' and title.lower() == 'преамбула'):
                items.append({'kind': 'heading', 'text': title})

            for paragraph in clause.findall('.//paragraph'):
                paragraph_text = ''.join(paragraph.itertext()).strip()
                if paragraph_text:
                    if not title_added and cls._looks_like_document_title(paragraph_text):
                        items.append({'kind': 'title', 'text': paragraph_text})
                        title_added = True
                        continue
                    items.append({'kind': 'paragraph', 'text': paragraph_text})

        for table in root.findall('./tables/table'):
            rows: List[List[str]] = []
            has_meaningful_cell = False
            for row in table.findall('./row'):
                row_values: List[str] = []
                for cell in row.findall('./cell'):
                    cell_text = ''.join(cell.itertext()).strip()
                    row_values.append(cell_text)
                    if cell_text:
                        has_meaningful_cell = True
                if row_values:
                    rows.append(row_values)
            if has_meaningful_cell:
                items.append({'kind': 'table', 'rows': rows})

        if not items:
            for text in root.xpath('.//text()'):
                clean = str(text).strip()
                if clean:
                    items.append({'kind': 'paragraph', 'text': clean})

        return items

    def _write_reconstructed_docx(self, contract: Contract, output_path: str, xml_content: str) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, Cm

        document = Document()

        for section in document.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)

        normal_style = document.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)

        content_items = self._extract_structured_content(xml_content)
        fallback_title = Path(contract.file_name or f"contract_{contract.id}").stem
        has_title = any(item['kind'] == 'title' for item in content_items)

        if not has_title and not content_items:
            title_paragraph = document.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run(fallback_title)
            title_run.bold = True
            title_run.font.size = Pt(14)
            title_run.font.name = 'Times New Roman'

        for item in content_items:
            if item['kind'] == 'table':
                rows = item.get('rows') or []
                if not rows:
                    continue
                max_cols = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'
                for row_idx, row_values in enumerate(rows):
                    for col_idx in range(max_cols):
                        table.rows[row_idx].cells[col_idx].text = row_values[col_idx] if col_idx < len(row_values) else ''
                continue

            text = item.get('text', '').strip()
            if not text:
                continue

            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = 'Times New Roman'

            if item['kind'] == 'title':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
                run.font.size = Pt(14)
                paragraph.paragraph_format.space_after = Pt(12)
            elif item['kind'] == 'heading':
                run.bold = True
                run.font.size = Pt(12)
                paragraph.paragraph_format.space_before = Pt(10)
                paragraph.paragraph_format.space_after = Pt(4)
            else:
                run.font.size = Pt(12)
                paragraph.paragraph_format.first_line_indent = Cm(1.25)
                paragraph.paragraph_format.space_after = Pt(4)

        document.save(output_path)

    def _write_reconstructed_pdf(self, contract: Contract, output_path: str, xml_content: str) -> None:
        from xml.sax.saxutils import escape
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as PDFTable, TableStyle
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError as exc:
            raise ValueError(
                "Конверсия в PDF временно недоступна: в окружении отсутствует библиотека reportlab"
            ) from exc

        styles = getSampleStyleSheet()
        base_font = 'Helvetica'
        font_candidates = [
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/Library/Fonts/Arial Unicode.ttf',
            '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
        ]
        for font_path in font_candidates:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont('ContractAIExport', font_path))
                    base_font = 'ContractAIExport'
                    break
                except Exception:
                    continue

        title_style = ParagraphStyle(
            'ContractExportTitle',
            parent=styles['Title'],
            fontName=base_font,
            fontSize=15,
            leading=18,
            spaceAfter=14,
        )
        heading_style = ParagraphStyle(
            'ContractExportHeading',
            parent=styles['Heading2'],
            fontName=base_font,
            fontSize=12,
            leading=15,
            spaceBefore=10,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            'ContractExportBody',
            parent=styles['BodyText'],
            fontName=base_font,
            fontSize=10,
            leading=14,
            spaceAfter=6,
        )

        def _safe_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
            return Paragraph(escape(text).replace('\n', '<br/>'), style)

        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=2.5 * cm,
            rightMargin=1.5 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )
        story: List[Any] = []
        content_items = self._extract_structured_content(xml_content)
        fallback_title = Path(contract.file_name or f"contract_{contract.id}").stem
        has_title = any(item['kind'] == 'title' for item in content_items)

        if not has_title and not content_items:
            story.append(_safe_paragraph(fallback_title, title_style))
            story.append(Spacer(1, 6))

        for item in content_items:
            if item['kind'] == 'table':
                rows = item.get('rows') or []
                if not rows:
                    continue
                max_cols = max(len(row) for row in rows)
                normalized_rows = []
                for row in rows:
                    normalized_rows.append([
                        _safe_paragraph((row[col_idx] if col_idx < len(row) else '').strip() or ' ', body_style)
                        for col_idx in range(max_cols)
                    ])
                pdf_table = PDFTable(normalized_rows, repeatRows=1)
                pdf_table.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(pdf_table)
                story.append(Spacer(1, 8))
                continue

            text = item.get('text', '').strip()
            if not text:
                continue
            if item['kind'] == 'title':
                story.append(_safe_paragraph(text, title_style))
            elif item['kind'] == 'heading':
                story.append(_safe_paragraph(text, heading_style))
            else:
                story.append(_safe_paragraph(text, body_style))

        if not story:
            story.append(_safe_paragraph(fallback_title, title_style))

        doc.build(story)

    @staticmethod
    def _xml_to_plain_text(xml_content: str) -> str:
        root = parse_xml_safely(xml_content)
        lines: List[str] = []

        for clause in root.findall('.//clause'):
            title = (clause.findtext('title', '') or '').strip()
            if title:
                lines.append(title)
            for paragraph in clause.findall('.//paragraph'):
                text = (paragraph.text or '').strip()
                if text:
                    lines.append(text)

        if not lines:
            for text in root.xpath('.//text()'):
                clean = str(text).strip()
                if clean:
                    lines.append(clean)

        return '\n'.join(lines)

    def _log_export(
        self,
        contract_id: str,
        export_format: str,
        file_paths: Dict[str, str],
        user_id: Optional[str]
    ) -> Optional[ExportLog]:
        """Log export to database"""
        try:
            successful_paths = {fmt: path for fmt, path in file_paths.items() if path}
            if not successful_paths:
                return None

            export_log = ExportLog(
                contract_id=contract_id,
                exported_by=user_id,
                export_type='quick_export',
                meta_info={
                    'requested_format': export_format,
                    'file_paths': successful_paths,
                    'file_sizes': {
                        fmt: os.path.getsize(path) if path and os.path.exists(path) else 0
                        for fmt, path in successful_paths.items()
                    },
                },
            )

            self.db.add(export_log)
            self.db.commit()
            self.db.refresh(export_log)

            return export_log

        except Exception as e:
            logger.error(f"Failed to log export: {e}")
            return None


__all__ = ["QuickExportAgent"]
